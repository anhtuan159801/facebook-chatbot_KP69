"""
Enhanced SCRAPER THỦ TỤC HÀNH CHÍNH - TẤT CẢ BỘ/NGÀNH
=======================================================
✅ Chọn được bất kỳ Bộ/Ngành nào
✅ Tải file hướng dẫn (.doc/.docx)
✅ Mỗi thủ tục một folder riêng
✅ Ghi rõ tên thủ tục trong danh sách
✅ Thêm link thủ tục vào cuối file Word
✅ Enhanced RAG integration with better content processing
✅ Improved chunking and embedding capabilities
✅ Better error handling and retry mechanisms
"""

from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.support.ui import Select
from selenium.common.exceptions import TimeoutException, NoSuchElementException
import time
import os
import re
import json
from urllib.parse import urljoin
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

import subprocess
import platform
import requests
from concurrent.futures import ThreadPoolExecutor, as_completed
import threading
import hashlib

# Import enhanced components
import sys
import os
from io import BytesIO

# Load environment variables from .env file
try:
    from dotenv import load_dotenv
    load_dotenv()
except ImportError:
    # If python-dotenv is not installed, continue without it
    pass

# Try to import each component with proper error handling
try:
    # Try relative imports first (when used as module)
    from .clean_vector_storage import VectorStorage
except (ImportError, ValueError, SystemError):
    # If relative import fails, try absolute import (when run as script)
    try:
        from clean_vector_storage import VectorStorage
    except ImportError:
        VectorStorage = None

try:
    from .embeddings_helper import get_embeddings_with_details, is_embeddings_available
except (ImportError, ValueError, SystemError):
    try:
        from embeddings_helper import get_embeddings_with_details, is_embeddings_available
    except ImportError:
        def get_embeddings_with_details(*args, **kwargs):
            return None
        is_embeddings_available = lambda: False

try:
    from .content_processor import ContentProcessor, ContentExtractionResult, ChunkingResult
except (ImportError, ValueError, SystemError):
    try:
        from content_processor import ContentProcessor, ContentExtractionResult, ChunkingResult
    except ImportError:
        ContentProcessor = None
        # Define placeholder classes
        class ContentExtractionResult:
            def __init__(self, content="", metadata=None, success=False, error=None):
                self.content = content
                self.metadata = metadata or {}
                self.success = success
                self.error = error

        class ChunkingResult:
            def __init__(self, chunks=None, metadata=None, success=False, error=None):
                self.chunks = chunks or []
                self.metadata = metadata or {}
                self.success = success
                self.error = error

# Danh sách các Bộ/Ngành
MINISTRIES = {
    "1": "Bộ Công an",
    "2": "Bộ Công thương",
    "3": "Bộ Dân tộc và Tôn giáo",
    "4": "Bộ Giáo dục và Đào tạo",
    "5": "Bộ Khoa học và Công nghệ",
    "6": "Bộ Ngoại giao",
    "7": "Bộ Nội vụ",
    "8": "Bộ Nông nghiệp và Môi trường",
    "9": "Bộ Quốc phòng",
    "10": "Bộ Tài chính",
    "11": "Bộ Tư pháp",
    "12": "Bộ Văn hóa, Thể thao và Du lịch",
    "13": "Bộ Y tế",
    "14": "Ngân hàng Chính sách xã hội",
    "15": "Ngân hàng Nhà nước Việt Nam",
    "16": "Ngân hàng phát triển Việt Nam",
    "17": "Thanh tra Chính phủ",
    "18": "Tòa án nhân dân",
    "19": "Tập đoàn Điện lực Việt Nam",
    "20": "Văn phòng Chính phủ"
}

import logging

class EnhancedErrorTracker:
    """Track and categorize errors for better debugging"""

    def __init__(self):
        self.errors = {
            'timeout': [],
            'conversion': [],
            'permission': [],
            'path_length': [],
            'missing_element': [],
            'other': []
        }
        self.lock = threading.Lock()

    def add_error(self, category, procedure_code, error_msg):
        """Add error to appropriate category"""
        with self.lock:
            if category in self.errors:
                self.errors[category].append({
                    'code': procedure_code,
                    'message': error_msg,
                    'timestamp': time.strftime('%Y-%m-%d %H:%M:%S')
                })
            else:
                self.errors['other'].append({
                    'code': procedure_code,
                    'message': error_msg,
                    'timestamp': time.strftime('%Y-%m-%d %H:%M:%S')
                })

    def get_summary(self):
        """Get error summary"""
        with self.lock:
            summary = {}
            for category, errors in self.errors.items():
                if errors:
                    summary[category] = len(errors)
            return summary

    def save_error_report(self, filepath):
        """Save detailed error report to file"""
        with self.lock:
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write("=" * 80 + "\n")
                f.write("ERROR REPORT\n")
                f.write("=" * 80 + "\n\n")

                for category, errors in self.errors.items():
                    if errors:
                        f.write(f"\n{category.upper()} ERRORS ({len(errors)}):\n")
                        f.write("-" * 80 + "\n")
                        for err in errors:
                            f.write(f"\nProcedure: {err['code']}\n")
                            f.write(f"Time: {err['timestamp']}\n")
                            f.write(f"Message: {err['message']}\n")
                            f.write("-" * 40 + "\n")


class EnhancedProcedureScraper:
    """Enhanced class to scrape administrative procedures with better RAG integration"""

    def __init__(self, download_dir='downloads', max_workers=8, headless=True):
        self.download_dir = os.path.abspath(download_dir)
        os.makedirs(self.download_dir, exist_ok=True)

        self.max_workers = max_workers
        self.headless = headless
        self.stats = {'success': 0, 'failed': 0, 'cached': 0}
        self.stats_lock = threading.Lock()

        # Cache
        self.cache_file = 'cache_ministries.json'
        self.cache = self._load_cache()
        self.cache_lock = threading.Lock()

        # Content processor for enhanced processing
        self.content_processor = ContentProcessor() if ContentProcessor else None

        # Main driver cho việc browse danh sách
        self.driver = self._create_driver()
        self.wait = WebDriverWait(self.driver, 15)

        # Session requests
        self.session = requests.Session()
        self.session.headers.update({
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
        })

        # FIX: Add error tracker
        self.error_tracker = EnhancedErrorTracker()

        # FIX: Setup enhanced logging
        self._setup_logging()

        print(f"Enhanced scraper initialized successfully")
        print(f"Download directory: {self.download_dir}")
        print(f"Headless mode: {self.headless}")
        print(f"Max workers: {self.max_workers}\n")

    def _setup_logging(self):
        """Setup comprehensive logging"""
        import logging
        from logging.handlers import RotatingFileHandler

        # Create logs directory
        log_dir = 'logs'
        os.makedirs(log_dir, exist_ok=True)

        # Setup main logger
        self.logger = logging.getLogger('ProcedureScraper')
        self.logger.setLevel(logging.DEBUG)

        # File handler
        log_file = os.path.join(log_dir, f'scraper_{time.strftime("%Y%m%d")}.log')
        file_handler = RotatingFileHandler(
            log_file,
            maxBytes=10*1024*1024,  # 10MB
            backupCount=5,
            encoding='utf-8'
        )
        file_handler.setLevel(logging.DEBUG)

        # Console handler
        console_handler = logging.StreamHandler()
        console_handler.setLevel(logging.INFO)

        # Formatter
        formatter = logging.Formatter(
            '%(asctime)s - %(name)s - %(levelname)s - %(message)s',
            datefmt='%Y-%m-%d %H:%M:%S'
        )
        file_handler.setFormatter(formatter)
        console_handler.setFormatter(formatter)

        self.logger.addHandler(file_handler)
        self.logger.addHandler(console_handler)

    def _create_driver(self):
        """Create Chrome driver with enhanced stability and error handling"""
        chrome_options = Options()

        if self.headless:
            chrome_options.add_argument('--headless=new')

        # Core stability options
        chrome_options.add_argument('--no-sandbox')
        chrome_options.add_argument('--disable-dev-shm-usage')
        chrome_options.add_argument('--disable-gpu')
        chrome_options.add_argument('--window-size=1920,1080')
        chrome_options.add_argument('--disable-blink-features=AutomationControlled')

        # Additional stability options
        chrome_options.add_argument('--disable-web-security')
        chrome_options.add_argument('--allow-running-insecure-content')
        chrome_options.add_argument('--disable-features=VizDisplayCompositor')

        # Performance and reliability options
        chrome_options.add_argument('--disable-extensions')
        chrome_options.add_argument('--disable-plugins')
        chrome_options.add_argument('--blink-settings=imagesEnabled=true')  # Keep images for better page rendering
        chrome_options.add_argument('--no-first-run')
        chrome_options.add_argument('--disable-features=TranslateUI')
        chrome_options.add_argument('--disable-ipc-flooding-protection')
        chrome_options.add_argument('--disable-backgrounding-occluded-windows')
        chrome_options.add_argument('--disable-renderer-backgrounding')

        # FIX: Disable Google API registration to avoid DEPRECATED_ENDPOINT error
        chrome_options.add_argument('--disable-sync')
        chrome_options.add_argument('--disable-background-networking')
        chrome_options.add_experimental_option('excludeSwitches', ['enable-automation', 'enable-logging'])
        chrome_options.add_experimental_option('useAutomationExtension', False)

        # FIX: Suppress console errors
        chrome_options.add_argument('--log-level=3')  # Suppress INFO, WARNING
        chrome_options.add_argument('--silent')

        # Suppress ChromeDriver logs
        import os
        os.environ['WDM_LOG'] = '0'  # Disable webdriver-manager logs

        try:
            # Create service with reduced logging
            from selenium.webdriver.chrome.service import Service
            service = Service(log_path='NUL' if platform.system() == 'Windows' else '/dev/null')

            driver = webdriver.Chrome(service=service, options=chrome_options)

            # FIX: Enhanced timeouts for stability
            driver.set_page_load_timeout(60)  # Increased from 30 to 60 seconds
            driver.set_script_timeout(30)     # Increased from 20 to 30 seconds
            driver.implicitly_wait(15)        # Increased from 8 to 15 seconds

            return driver

        except Exception as e:
            self.logger.error(f"Failed to create driver: {e}")
            raise

    def _load_cache(self):
        """Load cache from file"""
        if os.path.exists(self.cache_file):
            try:
                with open(self.cache_file, 'r', encoding='utf-8') as f:
                    return json.load(f)
            except json.JSONDecodeError as e:
                print(f"⚠️ Warning - Lỗi đọc file cache '{self.cache_file}': {e}. Tạo cache mới.")
                return {}
            except Exception as e:
                print(f"⚠️ Warning - Lỗi không xác định khi tải cache '{self.cache_file}': {e}. Tạo cache mới.")
                return {}
        return {}

    def _save_cache(self):
        """Save cache to file"""
        try:
            with open(self.cache_file, 'w', encoding='utf-8') as f:
                json.dump(self.cache, f, ensure_ascii=False, indent=2)
        except Exception as e:
            print(f"⚠️ Warning - Không thể lưu cache: {e}")

    def setup_advanced_search(self, ministry_name):
        """
        Setup advanced search for specific ministry
        """
        try:
            # BƯỚC 1: Click "Tìm kiếm nâng cao"
            print(f"🔍 Bước 1: Click 'Tìm kiếm nâng cao'...", end=" ")
            try:
                adv_search = self.wait.until(
                    EC.element_to_be_clickable((By.CSS_SELECTOR, "div.adv"))
                )
                adv_search.click()
                # Wait for the select2 dropdown to become visible after clicking advanced search
                self.wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, "span.select2-selection__rendered[id='select2-select-implementation-agency-container']")))
                print("✅")
            except TimeoutException:
                print("⚠️ (Có thể đã mở sẵn hoặc selector đã thay đổi)")
            except Exception as e:
                print(f"❌ Lỗi khi click tìm kiếm nâng cao: {e}")
                return False

            # BƯỚC 2: Choose Ministry from dropdown
            print(f"🏢 Bước 2: Chọn '{ministry_name}'...", end=" ")

            # Click into the Select2 dropdown
            select2_container = self.wait.until(
                EC.element_to_be_clickable((
                    By.CSS_SELECTOR,
                    "span.select2-selection__rendered[id='select2-select-implementation-agency-container']"
                ))
            )
            select2_container.click()

            # Wait for search box or options to appear
            try:
                # Try to find and click the option directly
                option = self.wait.until(
                    EC.element_to_be_clickable((
                        By.XPATH,
                        f"//li[contains(@class, 'select2-results__option') and contains(text(), '{ministry_name}')]"
                    ))
                )
                option.click()
                print("✅")
            except TimeoutException:
                # Fallback: Use search box if direct click fails
                try:
                    search_box = self.wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "input.select2-search__field")))
                    search_box.send_keys(ministry_name)
                    option = self.wait.until(
                        EC.element_to_be_clickable((
                            By.XPATH,
                            "//li[contains(@class, 'select2-results__option')]"
                        ))
                    )
                    option.click()
                    print("✅")
                except Exception as e:
                    print(f"❌ Không chọn được: {e}")
                    return False

            # Wait for the select2 dropdown to close or become invisible
            self.wait.until(EC.invisibility_of_element_located((By.CSS_SELECTOR, "input.select2-search__field")))

            # BƯỚC 3: Click search button
            print("🔎 Bước 3: Click nút 'Tìm kiếm'...", end=" ")
            search_btn = self.wait.until(
                EC.element_to_be_clickable((By.ID, "btn-search"))
            )
            search_btn.click()
            # Wait for the search results table to be present and contain data (or a loading indicator to disappear)
            self.wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, "table tbody tr td a[href*='ma_thu_tuc=']")))
            print("✅")

            # BƯỚC 4: Choose 50 rows per page
            print("📊 Bước 4: Chọn 50 hàng/trang...", end=" ")
            select_el = self.wait.until(
                EC.presence_of_element_located((By.ID, "paginationRecsPerPage"))
            )
            Select(select_el).select_by_value("50")
            # Wait for the table content to refresh after changing items per page
            self.wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, "table tbody tr td a[href*='ma_thu_tuc=']")))
            print("✅")

            print("✅ Setup hoàn tất!\n")
            return True

        except Exception as e:
            print(f"\n❌ Lỗi setup: {e}")
            import traceback
            traceback.print_exc()
            return False

    def get_current_page_number(self):
        """Get current page number"""
        try:
            active = self.driver.find_element(By.CSS_SELECTOR, "li.active[jp-role='page']")
            return int(active.get_attribute('jp-data'))
        except NoSuchElementException:
            return 1 # Assume page 1 if active page element not found
        except ValueError:
            print(f"   ⚠️ Could not parse current page number, defaulting to 1.")
            return 1
        except Exception as e:
            print(f"   ❌ Error getting current page number: {e}, defaulting to 1.")
            return 1

    def get_total_pages(self):
        """Get total number of pages"""
        try:
            last_btn = self.driver.find_element(By.CSS_SELECTOR, "li.last[jp-role='last']")
            total = int(last_btn.get_attribute('jp-data'))
            return total
        except NoSuchElementException:
            try:
                page_nums = self.driver.find_elements(By.CSS_SELECTOR, "li[jp-role='page']")
                if page_nums:
                    return max([int(p.get_attribute('jp-data')) for p in page_nums])
            except (NoSuchElementException, ValueError) as e:
                print(f"   ⚠️ Could not determine total pages from page numbers: {e}, defaulting to 1.")
            return 1 # Default to 1 if last page button or page numbers are not found
        except ValueError:
            print(f"   ⚠️ Could not parse total pages number, defaulting to 1.")
            return 1
        except Exception as e:
            print(f"   ❌ Error getting total pages: {e}, defaulting to 1.")
            return 1

    def go_to_next_page(self):
        """Go to the next page"""
        try:
            current_page = self.get_current_page_number()
            print(f"   📄 Đang ở trang {current_page}, chuyển sang trang {current_page + 1}...", end=" ")

            next_btn = self.wait.until(
                EC.element_to_be_clickable((By.CSS_SELECTOR, "li.next:not(.disabled) a"))
            )
            next_btn.click()

            # Wait for page number to change
            try:
                self.wait.until(EC.text_to_be_present_in_element_attribute(
                    (By.CSS_SELECTOR, "li.active[jp-role='page']"), 'jp-data', str(current_page + 1)
                ))
                print(f"✅ Đã chuyển sang trang {current_page + 1}")
                return True
            except TimeoutException:
                print(f"⚠️ Timeout when changing page")
                return False

        except Exception as e:
            print(f"❌ Lỗi: {e}")
            return False

    def extract_procedures_from_page(self):
        """Extract procedures from current page"""
        procedures = []
        seen_ids = set()

        try:
            # Wait for at least one procedure link to be present
            self.wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, "table tbody tr td a[href*='ma_thu_tuc=']")))

            links = self.driver.find_elements(
                By.CSS_SELECTOR,
                "table tbody tr td a[href*='ma_thu_tuc=']"
            )

            print(f"   🔍 Tìm thấy {len(links)} links")

            for link in links:
                try:
                    href = link.get_attribute('href')
                    if not href:
                        print(f"   ⚠️ Link trống, bỏ qua.")
                        continue

                    match = re.search(r'ma_thu_tuc=(\d+)', href)
                    if not match:
                        print(f"   ⚠️ Không tìm thấy ma_thu_tuc trong link: {href}, bỏ qua.")
                        continue

                    proc_id = match.group(1)

                    if proc_id in seen_ids:
                        print(f"   ⚠️ Thủ tục trùng lặp ID: {proc_id}, bỏ qua.")
                        continue
                    seen_ids.add(proc_id)

                    # Get display code
                    code = f"TTHC-{proc_id}" # Default value
                    try:
                        code_span = link.find_element(By.CSS_SELECTOR, "span.link.thick")
                        code = code_span.text.strip()
                    except NoSuchElementException:
                        pass # Use default code if thick span not found

                    # Get procedure name (all text except code)
                    full_text = link.text.strip()
                    title = full_text.replace(code, '').strip()
                    if not title:
                        title = f"Thủ tục {proc_id}"
                        print(f"   ⚠️ Tên thủ tục trống, dùng mặc định: {title}")

                    procedures.append({
                        'id': proc_id,
                        'code': code,
                        'title': title,
                        'url': href
                    })

                except Exception as link_e:
                    print(f"   ❌ Lỗi khi xử lý link thủ tục: {link_e}")
                    continue

        except TimeoutException:
            print(f"   ⚠️ Timeout when waiting for procedure table. No procedures on this page.")
        except Exception as e:
            print(f"   ⚠️ General error when extracting procedures: {e}")

        return procedures

    def get_all_procedures(self, url, ministry_name):
        """Get all procedures for a ministry"""
        print("="*70)
        print(f"🚀 BẮT ĐẦU LẤY DANH SÁCH THỦ TỤC - {ministry_name.upper()}")
        print("="*70)

        self.driver.get(url)
        try:
            # Wait for key element to ensure page is loaded
            self.wait.until(EC.presence_of_element_located((By.ID, "btn-search")))
        except TimeoutException:
            print("   ⚠️ Search button not loaded within timeout, trying to continue...")
            # Check if page at least loaded partially
            try:
                # Look for any element that indicates the page loaded
                self.wait.until(EC.presence_of_element_located((By.TAG_NAME, "body")))
                print("   ℹ️  Page body loaded, continuing...")
            except TimeoutException:
                print("   ❌ Page failed to load properly")
                return []

        if not self.setup_advanced_search(ministry_name):
            print("❌ Không thể setup tìm kiếm!")
            return []

        all_procedures = []
        total_pages = self.get_total_pages()
        current = self.get_current_page_number()

        print(f"📊 Tổng số trang: {total_pages}")
        print(f"📊 Trang hiện tại: {current}\n")

        for page_num in range(1, total_pages + 1):
            current_page = self.get_current_page_number()
            print(f"📄 Trang {current_page}/{total_pages}")

            procs = self.extract_procedures_from_page()

            if not procs:
                print("   ⚠️ Không lấy được thủ tục nào!")
                break

            all_procedures.extend(procs)
            print(f"   ✅ Lấy được {len(procs)} thủ tục | Tổng: {len(all_procedures)}\n")

            if page_num < total_pages:
                if not self.go_to_next_page():
                    print("   ⚠️ Dừng lại vì không thể chuyển trang\n")
                    break

        # Remove duplicates
        unique_dict = {}
        for p in all_procedures:
            if p['id'] not in unique_dict:
                unique_dict[p['id']] = p

        unique_procs = list(unique_dict.values())

        print(f"✅ TỔNG CỘNG: {len(unique_procs)} thủ tục duy nhất")
        if len(all_procedures) > len(unique_procs):
            print(f"   (Đã loại bỏ {len(all_procedures) - len(unique_procs)} duplicate)\n")

        return unique_procs

    def _convert_doc_to_docx(self, doc_path):
        """
        Enhanced conversion with better fallback handling
        """
        # If already .docx, return as is
        if doc_path.endswith('.docx'):
            return doc_path

        docx_path = doc_path.replace('.doc', '.docx')

        # If .docx exists, cleanup and return
        if os.path.exists(docx_path):
            try:
                if doc_path != docx_path and os.path.exists(doc_path):
                    os.remove(doc_path)
            except Exception as e:
                self.logger.warning(f"Cannot remove .doc: {e}")
            return docx_path

        # Try conversion methods in order
        conversion_methods = [
            ('Win32COM', self._convert_with_win32com),
            ('LibreOffice', self._convert_with_libreoffice),
            ('Pypandoc', self._convert_with_pypandoc),
        ]

        for method_name, method in conversion_methods:
            try:
                if method(doc_path, docx_path):
                    # Success
                    if os.path.exists(docx_path) and os.path.exists(doc_path):
                        os.remove(doc_path)
                        print(f"   ✅ Converted with {method_name}: {os.path.basename(doc_path)} → .docx")
                    return docx_path
            except Exception as e:
                self.logger.debug(f"{method_name} conversion failed: {e}")
                continue

        # FIX: All conversion methods failed
        print(f"   ⚠️ All conversion methods failed for {os.path.basename(doc_path)}")

        # FIX: Try to at least read the content for processing
        content = self._fallback_read_doc_content(doc_path)
        if content:
            print(f"   ℹ️  Keeping .doc file (conversion failed but content readable)")
            return doc_path  # Return .doc file, we can still process it
        else:
            print(f"   ❌ Cannot convert or read {os.path.basename(doc_path)}")
            # FIX: Don't delete the file, keep for manual review
            print(f"   ℹ️  File kept for manual review: {doc_path}")
            return doc_path

    def _get_safe_path(self, filepath, max_length=200):
        """
        Get a safe path that works with Windows limitations

        Windows has 260 character path limit (MAX_PATH)
        We need to keep it shorter to avoid issues
        """
        import os

        # Get absolute path
        abs_path = os.path.abspath(filepath)

        # Check if path is too long
        if len(abs_path) > max_length:
            # Try to shorten by using shorter directory names
            directory = os.path.dirname(abs_path)
            filename = os.path.basename(abs_path)

            # Create a shorter directory structure
            parts = directory.split(os.sep)

            # Keep drive and shorten middle parts
            drive = parts[0] if len(parts) > 0 else ""
            short_parts = [drive]

            for part in parts[1:]:
                # Shorten each directory name to max 30 chars
                if len(part) > 30:
                    short_part = part[:27] + "..."
                else:
                    short_part = part
                short_parts.append(short_part)

            short_dir = os.sep.join(short_parts)
            short_path = os.path.join(short_dir, filename)

            # If still too long, use temp directory
            if len(short_path) > max_length:
                import tempfile
                temp_dir = tempfile.gettempdir()
                short_path = os.path.join(temp_dir, "doc_convert", filename)
                os.makedirs(os.path.dirname(short_path), exist_ok=True)

            return short_path

        return abs_path


    def _convert_with_win32com(self, doc_path, docx_path):
        """
        Enhanced Win32COM conversion with path and permission handling
        """
        if platform.system() != 'Windows':
            return False

        try:
            import win32com.client
            import pythoncom

            # FIX: Get safe paths (short paths)
            safe_doc_path = self._get_safe_path(doc_path)
            safe_docx_path = self._get_safe_path(docx_path)

            # FIX: If paths are different, copy file to safe location
            if safe_doc_path != doc_path:
                import shutil
                os.makedirs(os.path.dirname(safe_doc_path), exist_ok=True)
                shutil.copy2(doc_path, safe_doc_path)
                print(f"   Using shorter path for conversion: {safe_doc_path}")

            # Initialize COM
            pythoncom.CoInitialize()

            word = None
            doc = None
            conversion_success = False

            try:
                # FIX: Check write permissions
                test_file = safe_docx_path + ".test"
                try:
                    with open(test_file, 'w') as f:
                        f.write("test")
                    os.remove(test_file)
                except PermissionError:
                    print(f"   ⚠️ No write permission for: {os.path.dirname(safe_docx_path)}")
                    return False
                except OSError as e:
                    print(f"   ⚠️ Cannot write to directory: {e}")
                    return False

                # FIX: Create Word application with error handling
                try:
                    word = win32com.client.DispatchEx("Word.Application")
                except Exception as e:
                    print(f"   ⚠️ Cannot create Word application: {e}")
                    return False

                word.Visible = False
                word.DisplayAlerts = 0  # wdAlertsNone

                # FIX: Open document with enhanced error handling
                try:
                    doc = word.Documents.Open(
                        os.path.abspath(safe_doc_path),
                        ReadOnly=True,
                        Visible=False,
                        ConfirmConversions=False,
                        AddToRecentFiles=False,
                        Revert=False
                    )
                except Exception as e:
                    print(f"   ⚠️ Cannot open document: {e}")
                    return False

                # FIX: Save with enhanced error handling
                try:
                    # Ensure output directory exists
                    os.makedirs(os.path.dirname(safe_docx_path), exist_ok=True)

                    # Try SaveAs2 first (newer)
                    try:
                        doc.SaveAs2(
                            os.path.abspath(safe_docx_path),
                            FileFormat=16,  # wdFormatXMLDocument
                            AddToRecentFiles=False
                        )
                    except AttributeError:
                        # Fallback to SaveAs for older Word versions
                        doc.SaveAs(
                            os.path.abspath(safe_docx_path),
                            FileFormat=16
                        )

                    conversion_success = os.path.exists(safe_docx_path)

                    # FIX: If we used temp path, copy back to original location
                    if conversion_success and safe_docx_path != docx_path:
                        os.makedirs(os.path.dirname(docx_path), exist_ok=True)
                        shutil.copy2(safe_docx_path, docx_path)
                        conversion_success = os.path.exists(docx_path)
                        # Cleanup temp files
                        try:
                            os.remove(safe_docx_path)
                            if safe_doc_path != doc_path:
                                os.remove(safe_doc_path)
                        except:
                            pass

                except Exception as e:
                    error_msg = str(e)

                    # FIX: Handle specific error cases
                    if "write-protected" in error_msg.lower() or "disk" in error_msg.lower():
                        print(f"   ⚠️ Write protection or disk error: {error_msg[:100]}")
                        # Try alternate location
                        import tempfile
                        alt_path = os.path.join(tempfile.gettempdir(), os.path.basename(safe_docx_path))
                        try:
                            doc.SaveAs2(alt_path, FileFormat=16)
                            if os.path.exists(alt_path):
                                shutil.copy2(alt_path, docx_path)
                                os.remove(alt_path)
                                conversion_success = os.path.exists(docx_path)
                        except:
                            pass

                    if not conversion_success:
                        print(f"   ⚠️ Cannot save document: {error_msg[:100]}")
                        return False

                return conversion_success

            finally:
                # FIX: Enhanced cleanup
                if doc:
                    try:
                        doc.Close(SaveChanges=False)
                    except Exception as e:
                        print(f"   ⚠️ Error closing document: {e}")

                if word:
                    try:
                        word.Quit()
                    except Exception as e:
                        print(f"   ⚠️ Error quitting Word: {e}")

                    # FIX: Force cleanup COM objects
                    try:
                        del doc
                        del word
                    except:
                        pass

                # Uninitialize COM
                try:
                    pythoncom.CoUninitialize()
                except:
                    pass

        except ImportError:
            self.logger.debug("win32com not available")
            return False
        except Exception as e:
            self.logger.error(f"Win32COM conversion error: {e}")
            return False

    def _convert_with_libreoffice(self, doc_path, docx_path):
        """
        Convert using LibreOffice command line
        """
        try:
            # Find LibreOffice executable
            libreoffice_paths = {
                'Windows': [
                    r'C:\Program Files\LibreOffice\program\soffice.exe',
                    r'C:\Program Files (x86)\LibreOffice\program\soffice.exe',
                ],
                'Darwin': [  # macOS
                    '/Applications/LibreOffice.app/Contents/MacOS/soffice',
                ],
                'Linux': [
                    '/usr/bin/libreoffice',
                    '/usr/bin/soffice',
                ]
            }

            system = platform.system()
            soffice_path = None

            for path in libreoffice_paths.get(system, []):
                if os.path.exists(path):
                    soffice_path = path
                    break

            if not soffice_path:
                # Try to find in PATH
                import shutil
                soffice_path = shutil.which('soffice') or shutil.which('libreoffice')

            if not soffice_path:
                return False

            # Convert
            output_dir = os.path.dirname(doc_path)

            result = subprocess.run(
                [
                    soffice_path,
                    '--headless',
                    '--convert-to', 'docx',
                    '--outdir', output_dir,
                    doc_path
                ],
                capture_output=True,
                timeout=60,
                text=True
            )

            return result.returncode == 0 and os.path.exists(docx_path)

        except Exception as e:
            self.logger.error(f"LibreOffice conversion error: {e}")
            return False

    def _convert_with_pypandoc(self, doc_path, docx_path):
        """
        Convert using pypandoc with better error handling
        """
        try:
            import pypandoc

            # FIX: Check if pandoc is actually available
            try:
                pypandoc.get_pandoc_version()
            except OSError:
                # Pandoc not installed
                self.logger.debug("Pandoc binary not found")
                return False

            # Perform conversion
            pypandoc.convert_file(
                doc_path,
                'docx',
                outputfile=docx_path,
                extra_args=['--wrap=none']
            )

            return os.path.exists(docx_path)

        except ImportError:
            # FIX: Provide helpful installation message ONCE
            if not hasattr(self, '_pypandoc_warning_shown'):
                print("\n   ℹ️  Pypandoc not installed. To enable:")
                print("      pip install pypandoc")
                print("      Then install pandoc binary:")
                print("      - Windows: choco install pandoc")
                print("      - Mac: brew install pandoc")
                print("      - Linux: apt-get install pandoc\n")
                self._pypandoc_warning_shown = True
            return False
        except Exception as e:
            self.logger.error(f"Pypandoc conversion error: {e}")
            return False


    # FIX: Add completely fallback method using textract for reading only
    def _fallback_read_doc_content(self, doc_path):
        """
        Fallback to just read .doc content without conversion
        This is used when all conversion methods fail
        """
        try:
            import textract  # type: ignore
            content = textract.process(doc_path).decode('utf-8')
            return content
        except ImportError:
            return None
        except Exception as e:
            self.logger.error(f"Textract read error: {e}")
            return None

    def _sanitize_filename(self, filename):
        """Sanitize filename by removing invalid characters for Windows"""
        if not filename:
            return "unnamed_file"

        # Ensure it's a string
        filename = str(filename)

        # Remove invalid characters for Windows file names
        invalid_chars = '<>:"/\\|?*'
        for char in invalid_chars:
            filename = filename.replace(char, '_')

        # Also replace other problematic characters including ampersand, etc.
        problematic_chars = ['&', '%', '#', '@', '!', '+', '=', '~', '`', '$', '^', '{', '}', '[', ']', ';', ',', '(', ')']
        for char in problematic_chars:
            filename = filename.replace(char, '_')

        # Replace control characters and other potentially problematic characters
        filename = ''.join(c if ord(c) >= 32 and c not in '<>:"/\\|?*[]{}&%#@!+=~`$^;(),' else '_' for c in filename)

        # Also handle reserved names in Windows by adding underscore
        reserved_names = ['CON', 'PRN', 'AUX', 'NUL'] + \
                        [f'COM{i}' for i in range(1, 10)] + \
                        [f'LPT{i}' for i in range(1, 10)]

        # Check just the name part before extension
        name_part = filename.split('.')[0] if '.' in filename else filename
        if name_part.upper() in reserved_names:
            filename = f"{filename}_"

        # Remove any trailing periods or spaces (Windows doesn't allow these at the end)
        filename = filename.rstrip('. ')

        # Limit length more conservatively to prevent path too long errors
        # Leave room for the full path including directory structure
        if len(filename) > 80:  # More conservative limit
            filename = filename[:80]

        # If filename is empty after sanitization, provide a default
        if not filename or filename.replace('_', '').replace('.', '') == '':
            filename = "unnamed_file"

        return filename

    def _calculate_file_checksum(self, filepath):
        """Calculate MD5 checksum of a file"""
        hash_md5 = hashlib.md5()
        try:
            with open(filepath, "rb") as f:
                # Read file in chunks to handle large files efficiently
                for chunk in iter(lambda: f.read(4096), b""):
                    hash_md5.update(chunk)
            return hash_md5.hexdigest()
        except FileNotFoundError:
            print(f"   ⚠️ File not found for checksum calculation: {filepath}")
            return None
        except IOError as e:
            print(f"   ❌ Error reading file for checksum calculation {filepath}: {e}")
            return None
        except Exception as e:
            print(f"   ❌ Unexpected error during checksum calculation {filepath}: {e}")
            return None

    def extract_content_from_docx_path(self, docx_path: str) -> str:
        """
        Extract text content from a .doc or .docx file path with improved structure handling
        """
        try:
            content_parts = []

            if docx_path.lower().endswith('.docx'):
                # Handle .docx files using python-docx
                doc = Document(docx_path)

                # Extract text from paragraphs with better formatting
                for i, paragraph in enumerate(doc.paragraphs):
                    text = paragraph.text.strip()
                    if text:
                        content_parts.append(text)

                # Extract text from tables with proper structure
                for table_idx, table in enumerate(doc.tables):
                    content_parts.append(f"\n--- Bảng {table_idx + 1} ---")
                    for row_idx, row in enumerate(table.rows):
                        row_data = []
                        for cell_idx, cell in enumerate(row.cells):
                            cell_text = cell.text.strip()
                            if cell_text:
                                row_data.append(cell_text)

                        if row_data:
                            content_parts.append(" | ".join(row_data))
                    content_parts.append("--- Hết bảng ---\n")

                # Extract and format list items if any
                for paragraph in doc.paragraphs:
                    p_style = paragraph.style.name.lower() if paragraph.style.name else ""
                    if 'list' in p_style or 'bullet' in p_style or paragraph.text.strip().startswith(('•', '○', '-', '–')):
                        content_parts.append(f"• {paragraph.text.strip()}")

                # Extract text from tables with more detail
                for table_idx, table in enumerate(doc.tables):
                    content_parts.append(f"\n--- Chi tiết bảng {table_idx + 1} ---")
                    headers = []
                    for i, cell in enumerate(table.rows[0].cells if table.rows else []):
                        headers.append(cell.text.strip())

                    for row_idx, row in enumerate(table.rows[1:], start=1):  # Skip header row
                        row_data = []
                        for cell_idx, cell in enumerate(row.cells):
                            cell_text = cell.text.strip()
                            if headers and cell_idx < len(headers):
                                row_data.append(f"{headers[cell_idx]}: {cell_text}")
                            else:
                                row_data.append(cell_text)

                        if row_data:
                            content_parts.append(" | ".join(row_data))
                    content_parts.append("--- Hết chi tiết bảng ---\n")

            elif docx_path.lower().endswith('.doc'):
                # Handle .doc files using textract as fallback
                try:
                    import textract  # type: ignore
                    content = textract.process(docx_path).decode('utf-8')
                    content_parts.append(content)
                except ImportError:
                    # Only print to console if we can handle the encoding
                    try:
                        print(f"   Warning: textract not installed for .doc extraction: pip install textract")
                    except UnicodeEncodeError:
                        pass
                    # Provide a helpful message to the user
                    return f"Content extraction from .doc files requires textract. Install with: pip install textract\nFile: {docx_path}"
                except Exception as doc_error:
                    try:
                        print(f"   Error when extracting content from .doc file {docx_path}: {doc_error}")
                    except UnicodeEncodeError:
                        pass
                    # Try alternative method with antiword if available (mainly for Linux)
                    try:
                        import subprocess
                        result = subprocess.run(['antiword', docx_path], capture_output=True, text=True, check=True)
                        content_parts.append(result.stdout)
                    except (subprocess.CalledProcessError, FileNotFoundError):
                        try:
                            print("   Warning: antiword not available. Install with: sudo apt-get install antiword (Ubuntu/Debian)")
                        except UnicodeEncodeError:
                            pass
                        return f"Content extraction from .doc failed. Install textract or antiword.\nFile: {docx_path}"

            return '\n'.join(content_parts)
        except Exception as e:
            try:
                print(f"   Error when extracting content from {docx_path}: {e}")
            except UnicodeEncodeError:
                pass
            return ""

    def extract_structural_info_from_content(self, content: str) -> dict:
        """
        Extract structural information from content that matches typical procedure document patterns
        Enhanced version with better pattern recognition
        """
        import re

        try:
            extracted_info = {}

            # Extract procedure code if mentioned in the document
            # Look for patterns like "Mã số: [code]" or "Mã thủ tục: [code]"
            code_patterns = [
                r'Mã thủ tục[:\s]*([^\n\r]+)',
                r'Mã số thủ tục[:\s]*([^\n\r]+)',
                r'Mã số[:\s]*([^\n\r]+)',
                r'TTHC[-_\s]*([0-9.]+)'  # TTHC-12345 or TTHC 12345 or TTHC_12345
            ]

            for pattern in code_patterns:
                match = re.search(pattern, content, re.IGNORECASE)
                if match:
                    extracted_info['procedure_code'] = match.group(1).strip()
                    break

            # Extract procedure title
            title_patterns = [
                r'Tên thủ tục[:\s]*([^\n\r]+)',
                r'Tên đầy đủ[:\s]*([^\n\r]+)',
                r'THỦ TỤC[:\s]*([^\n\r]+)',
                r'([^:\n\r]{20,100})'  # Potential title from first substantial line
            ]

            for pattern in title_patterns:
                match = re.search(pattern, content, re.IGNORECASE)
                if match:
                    title = match.group(1).strip()
                    # Ensure it's actually a title, not just random text
                    if len(title) > 5 and len(title) < 200:
                        extracted_info['procedure_title'] = title
                        break

            # Extract processing time - more comprehensive patterns
            time_patterns = [
                r'Thời hạn giải quyết[:\s]*([^\n\r.]+)',
                r'Thời gian giải quyết[:\s]*([^\n\r.]+)',
                r'Thời hạn[:\s]*([^\n\r.]+)',
                r'(\d+\s*(ngày|ngay|day|working day|ngày làm việc|ngày làm việc|ngày công tác|business day))'
            ]

            for pattern in time_patterns:
                match = re.search(pattern, content, re.IGNORECASE)
                if match:
                    extracted_info['processing_time'] = match.group(1).strip()
                    break

            # Extract fee information
            fee_patterns = [
                r'Phí[:\s]*([^\n\r.]+)',
                r'Lệ phí[:\s]*([^\n\r.]+)',
                r'Phí, lệ phí[:\s]*([^\n\r.]+)',
                r'([0-9.,\s]+đồng|miễn phí|không thu phí|0 đồng|0đ|[0-9]+\.?[0-9]*\s*vnd)'
            ]

            for pattern in fee_patterns:
                match = re.search(pattern, content, re.IGNORECASE)
                if match:
                    extracted_info['fee'] = match.group(1).strip()
                    break

            # Extract responsible authority
            agency_patterns = [
                r'Cơ quan thực hiện[:\s]*([^\n\r.]+)',
                r'Cơ quan có thẩm quyền[:\s]*([^\n\r.]+)',
                r'Cơ quan[:\s]*([^\n\r.]+)',
                r'Đơn vị thực hiện[:\s]*([^\n\r.]+)'
            ]

            for pattern in agency_patterns:
                match = re.search(pattern, content, re.IGNORECASE)
                if match:
                    extracted_info['responsible_agency'] = match.group(1).strip()
                    break

            # Extract required documents - improved extraction
            docs_patterns = [
                r'(Thành phần hồ sơ|Hồ sơ bao gồm|Các loại giấy tờ|Giấy tờ cần nộp)[:\s]*([^\n\r]+(?:\n[^\n\r]+)*?)(?:\n[^:\n\r]{10,}|$)',
                r'(Bước|1\.|2\.|3\.|I\.|II\.|III\.)[^:\n\r]+(?:\n[^\n\r]+)*?hồ sơ[:\s]*([^\n\r]+(?:\n[^\n\r]+)*)'
            ]

            for pattern in docs_patterns:
                match = re.search(pattern, content, re.IGNORECASE)
                if match:
                    extracted_info['required_documents'] = match.group(2).strip() if match.lastindex >= 2 else match.group(1).strip()
                    break

            # Extract procedure steps - improved extraction
            steps_patterns = [
                r'(Trình tự thực hiện|Các bước thực hiện|Cách thức thực hiện|Thủ tục thực hiện)[:\s]*([^\n\r]+(?:\n[^\n\r]+)*?)(?:\n[^:\n\r]{10,}|$)',
                r'(Bước 1|Bước 2|Bước 3|Bước thứ nhất|Bước thứ hai|Bước thứ ba)[:\s]*([^\n\r]+(?:\n[^\n\r]+)*)'
            ]

            for pattern in steps_patterns:
                match = re.search(pattern, content, re.IGNORECASE)
                if match:
                    extracted_info['procedure_steps'] = match.group(2).strip() if match.lastindex >= 2 else match.group(1).strip()
                    break

            # Extract field/domain of the procedure
            field_patterns = [
                r'Lĩnh vực[:\s]*([^\n\r.]+)',
                r'Ngành[:\s]*([^\n\r.]+)',
                r'Chủ đề[:\s]*([^\n\r.]+)',
                r'([A-Z][a-zÀ-ỹ\s]+(?:[A-Z][a-zÀ-ỹ\s]*\s*)+) lĩnh vực'  # Pattern like "Đăng ký kinh doanh lĩnh vực"
            ]

            for pattern in field_patterns:
                match = re.search(pattern, content, re.IGNORECASE)
                if match:
                    extracted_info['procedure_field'] = match.group(1).strip()
                    break

            # Extract contact information if available
            contact_patterns = [
                r'(Địa chỉ|Address)[:\s]*([^\n\r.]+)',
                r'(Điện thoại|Phone|SĐT)[:\s]*([^\n\r.]+)',
                r'(Email|Thư điện tử)[:\s]*([^\n\r.]+)'
            ]

            for pattern in contact_patterns:
                matches = re.findall(pattern, content, re.IGNORECASE)
                if matches:
                    key = matches[0][0].strip()
                    value = matches[0][1].strip() if len(matches[0]) > 1 else ""
                    if 'contact' not in extracted_info:
                        extracted_info['contact'] = {}
                    extracted_info['contact'][key] = value

            return extracted_info

        except Exception as e:
            print(f"   ❌ Lỗi khi trích xuất thông tin cấu trúc từ nội dung: {e}")
            return {}

    def preprocess_content_for_embedding(self, content: str, metadata: dict = None) -> str:
        """
        Preprocess content to make it more suitable for embedding/semantic search
        This creates a structured representation that better captures the key information
        """
        try:
            if not content:
                return ""

            # Extract structural information if not provided
            if not metadata:
                metadata = self.extract_structural_info_from_content(content)

            structured_parts = []

            # Add key information at the beginning for better retrieval
            if metadata.get('procedure_title'):
                structured_parts.append(f"TIÊU ĐỀ THỦ TỤC: {metadata['procedure_title']}")

            if metadata.get('procedure_code'):
                structured_parts.append(f"MÃ THỦ TỤC: {metadata['procedure_code']}")

            if metadata.get('responsible_agency'):
                structured_parts.append(f"CƠ QUAN THỰC HIỆN: {metadata['responsible_agency']}")

            if metadata.get('processing_time'):
                structured_parts.append(f"THỜI HẠN GIẢI QUYẾT: {metadata['processing_time']}")

            if metadata.get('fee'):
                structured_parts.append(f"PHÍ LỆ PHÍ: {metadata['fee']}")

            if metadata.get('required_documents'):
                structured_parts.append(f"HỒ SƠ CẦN NỘP: {metadata['required_documents']}")

            if metadata.get('procedure_steps'):
                structured_parts.append(f"CÁC BƯỚC THỰC HIỆN: {metadata['procedure_steps']}")

            # Add the main content
            structured_parts.append(f"NỘI DUNG CHI TIẾT:")
            structured_parts.append(content)

            return "\n\n".join(structured_parts)
        except Exception as e:
            print(f"   ❌ Lỗi khi tiền xử lý nội dung để nhúng: {e}")
            # Fallback to original content if preprocessing fails
            return content

    def chunk_content_improved(self, content: str, procedure_metadata: dict = None,
                              target_words: int = 500, max_words: int = 800) -> list:
        """
        Improved chunking optimized for ~1000 word documents

        Args:
            content: Text content to chunk
            procedure_metadata: Metadata about the procedure
            target_words: Target words per chunk (default 500)
            max_words: Maximum words per chunk (default 800)
        """
        if not content:
            return []

        # Đếm tổng số từ
        total_words = self._count_vietnamese_words(content)

        # Nếu document ngắn (<= 800 từ), giữ nguyên 1 chunk duy nhất
        if total_words <= max_words:
            # Thêm metadata vào đầu nếu có
            if procedure_metadata:
                header = self._create_metadata_header(procedure_metadata)
                return [header + "\n\n" + content]
            return [content]

        # Document dài hơn: chia thành nhiều chunks
        chunks = []

        # Tạo header từ metadata
        header = ""
        if procedure_metadata:
            header = self._create_metadata_header(procedure_metadata) + "\n\n"

        # Split theo sections rõ ràng
        sections = self._split_into_sections(content)

        current_chunk = header
        current_words = self._count_vietnamese_words(header)

        for section in sections:
            section_words = self._count_vietnamese_words(section)

            # Nếu section này + chunk hiện tại vẫn < max_words
            if current_words + section_words <= max_words:
                current_chunk += section + "\n\n"
                current_words += section_words
            else:
                # Save chunk hiện tại (nếu không rỗng)
                if current_chunk.strip() and current_chunk.strip() != header.strip():
                    chunks.append(current_chunk.strip())

                # Bắt đầu chunk mới
                # Nếu section quá dài, chia nhỏ thêm
                if section_words > max_words:
                    sub_chunks = self._split_long_section(section, header, target_words, max_words)
                    chunks.extend(sub_chunks)
                    current_chunk = header
                    current_words = self._count_vietnamese_words(header)
                else:
                    current_chunk = header + section + "\n\n"
                    current_words = self._count_vietnamese_words(header) + section_words

        # Thêm chunk cuối cùng
        if current_chunk.strip() and current_chunk.strip() != header.strip():
            chunks.append(current_chunk.strip())

        # Filter chunks quá ngắn (< 50 words)
        min_words = 50
        filtered_chunks = [
            chunk for chunk in chunks
            if self._count_vietnamese_words(chunk) >= min_words
        ]

        return filtered_chunks

    def _create_metadata_header(self, metadata: dict) -> str:
        """Create a structured header from metadata"""
        header_parts = []

        if metadata.get('procedure_code'):
            header_parts.append(f"Mã thủ tục: {metadata['procedure_code']}")

        if metadata.get('procedure_title'):
            header_parts.append(f"Tên thủ tục: {metadata['procedure_title']}")

        if metadata.get('responsible_agency'):
            header_parts.append(f"Cơ quan: {metadata['responsible_agency']}")

        if metadata.get('processing_time'):
            header_parts.append(f"Thời hạn: {metadata['processing_time']}")

        if metadata.get('fee'):
            header_parts.append(f"Phí: {metadata['fee']}")

        return "\n".join(header_parts)

    def _split_into_sections(self, content: str) -> list:
        """
        Split content into logical sections based on Vietnamese admin doc structure
        """
        import re

        # Patterns for Vietnamese administrative document sections
        section_patterns = [
            r'^(Điều \d+[.:]\s*.+?)(?=^Điều \d+[.:]|\Z)',  # Điều 1, Điều 2
            r'^(Khoản \d+[.:]\s*.+?)(?=^Khoản \d+[.:]|\Z)',  # Khoản 1, Khoản 2
            r'^([IVX]+\.\s*.+?)(?=^[IVX]+\.|\Z)',  # I., II., III.
            r'^(\d+\.\s*.+?)(?=^\d+\.|\Z)',  # 1., 2., 3.
            r'^([a-z]\)\s*.+?)(?=^[a-z]\)|\Z)',  # a), b), c)
        ]

        sections = []

        # Thử split theo từng pattern
        for pattern in section_patterns:
            matches = re.finditer(pattern, content, re.MULTILINE | re.DOTALL)
            found_sections = [m.group(1).strip() for m in matches]

            if found_sections and len(found_sections) > 1:
                sections = found_sections
                break

        # Nếu không tìm được sections, split theo paragraphs
        if not sections:
            sections = [p.strip() for p in content.split('\n\n') if p.strip()]

        return sections

    def _split_long_section(self, section: str, header: str,
                           target_words: int, max_words: int) -> list:
        """
        Split a long section into smaller chunks
        """
        chunks = []

        # Split theo câu (dấu chấm, chấm hỏi, chấm than)
        import re
        sentences = re.split(r'([.!?]+\s+)', section)

        # Ghép lại các câu với dấu câu
        parts = []
        for i in range(0, len(sentences) - 1, 2):
            if i + 1 < len(sentences):
                parts.append(sentences[i] + sentences[i + 1])
            else:
                parts.append(sentences[i])

        current_chunk = header
        current_words = self._count_vietnamese_words(header)

        for part in parts:
            part_words = self._count_vietnamese_words(part)

            if current_words + part_words <= max_words:
                current_chunk += part
                current_words += part_words
            else:
                if current_chunk.strip() != header.strip():
                    chunks.append(current_chunk.strip())
                current_chunk = header + part
                current_words = self._count_vietnamese_words(header) + part_words

        # Chunk cuối cùng
        if current_chunk.strip() != header.strip():
            chunks.append(current_chunk.strip())

        return chunks

    def _break_long_section(self, section_content: str, chunk_size: int, header_str: str) -> list:
        """
        Break down a long section into smaller chunks
        """
        chunks = []

        # Split by sentences or newlines
        import re
        sentences = re.split(r'([.!?]+\s+|\n+)', section_content)

        # Reattach the sentence terminators
        parts = []
        for i, sent in enumerate(sentences):
            if i < len(sentences) - 1 and re.match(r'^[.!?]+\s*$', sentences[i+1]):
                parts.append(sent + sentences[i+1])
            elif not re.match(r'^[.!?]+\s*$', sent):
                parts.append(sent)

        current_chunk = header_str
        current_length = len(header_str)

        for part in parts:
            part = part.strip()
            if not part:
                continue

            if current_length + len(part) > chunk_size and len(current_chunk) > len(header_str):
                # Current chunk is full, save it and start a new one
                chunks.append(current_chunk.strip())
                current_chunk = header_str + part + "\n"
                current_length = len(header_str) + len(part) + 1
            else:
                current_chunk += part + "\n"
                current_length += len(part) + 1

        # Add the final chunk if it has content
        if current_chunk.strip() and current_chunk != header_str:
            chunks.append(current_chunk.strip())

        return chunks

    def process_document_for_storage(self, doc_path: str, procedure_info: dict):
        """
        Process a document (.doc or .docx) to extract full content for direct storage in knowledge base
        """
        try:
            if not self.content_processor:
                print("   ❌ Content processor not available")
                return None

            # Extract full content from the document using the enhanced content processor
            extraction_result = self.content_processor.extract_content_from_docx_path(doc_path)

            # Check if the result is a ContentExtractionResult object or a string
            if hasattr(extraction_result, 'success'):  # It's a ContentExtractionResult object
                if not extraction_result.success or not extraction_result.content.strip():
                    print(f"   ❌ No content extracted from {doc_path}")
                    return None
                full_content = extraction_result.content
            else:  # It's likely a string (older implementation)
                content_str = extraction_result if isinstance(extraction_result, str) else ""
                if not content_str or not content_str.strip():
                    print(f"   ❌ No content extracted from {doc_path}")
                    return None
                full_content = content_str

            # Extract structural information from content
            metadata = self.content_processor.extract_structural_info_from_content(full_content)

            # Update metadata with procedure info
            metadata.update(procedure_info)

            # Preprocess content for better storage
            processed_content = self.content_processor.preprocess_content_for_embedding(
                full_content,
                metadata
            )

            print(f"   ✅ Processed {doc_path}: Full content extracted (length: {len(processed_content)} chars)")
            return {
                'content': processed_content,
                'chunks': None,  # No chunks - storing full content
                'metadata': metadata
            }
        except Exception as e:
            print(f"   ❌ Error processing document for storage: {e}")
            import traceback
            traceback.print_exc()
            return None

    def store_documents_with_embeddings(self, ministry_name: str, source_url: str = None, upload_to_storage=True):
        """
        Process and store all downloaded documents to vector storage with embeddings
        Can optionally upload files to Supabase storage as well
        """
        if VectorStorage is None:
            print("   ❌ VectorStorage module not available. Please install required dependencies.")
            return False

        # Check if environment variables are set for Supabase
        import os
        supabase_url = os.getenv("SUPABASE_URL")
        supabase_key = os.getenv("SUPABASE_KEY")

        if not supabase_url or not supabase_key:
            print("   ❌ SUPABASE_URL and SUPABASE_KEY environment variables are required")
            return False

        try:
            # Initialize vector storage
            vector_store = VectorStorage(supabase_url=supabase_url, supabase_key=supabase_key)

            # Find all docx files in the ministry folder
            safe_ministry_name = self._sanitize_filename(ministry_name)
            ministry_folder = os.path.join(self.download_dir, safe_ministry_name)
            if not os.path.exists(ministry_folder):
                print(f"   ❌ Ministry folder does not exist: {ministry_folder}")
                return False

            doc_files = []
            for root, dirs, files in os.walk(ministry_folder):
                for file in files:
                    if file.lower().endswith(('.doc', '.docx')):
                        doc_files.append(os.path.join(root, file))

            if not doc_files:
                print(f"   ❌ No .doc or .docx files found in {ministry_folder}")
                return False

            print(f"   📁 Found {len(doc_files)} .doc/.docx files to process")

            # Process and store each document
            processed_count = 0
            failed_count = 0
            uploaded_count = 0

            for doc_file in doc_files:
                try:
                    # Extract procedure info from the file path or name
                    filename = os.path.basename(doc_file)
                    procedure_code = os.path.splitext(filename)[0]  # Get name without extension
                    procedure_title = f"Thủ tục {procedure_code}"

                    # Get the procedure URL from the list file if available
                    # Use the properly sanitized ministry name
                    safe_ministry_name = self._sanitize_filename(ministry_name)
                    list_file = os.path.join(ministry_folder, f'danh_sach_{safe_ministry_name}.txt')
                    actual_source_url = source_url
                    if os.path.exists(list_file):
                        with open(list_file, 'r', encoding='utf-8') as f:
                            content = f.read()
                            # Try to find the URL for this specific procedure in the list file
                            import re
                            # Look for the procedure code in the list file and extract the URL
                            pattern = rf'\[{re.escape(procedure_code)}\].*?\n\s*URL:\s*(https?://[^\n]+)'
                            match = re.search(pattern, content, re.DOTALL | re.IGNORECASE)
                            if match:
                                actual_source_url = match.group(1).strip()

                    # Upload the file to Supabase storage if requested
                    file_storage_url = None
                    if upload_to_storage:
                        file_storage_url = self.upload_file_to_supabase_storage(
                            file_path=doc_file,
                            bucket_name="government-documents",  # Using a dedicated bucket for gov documents
                            folder_path=f"{safe_ministry_name}/procedures"
                        )

                    # Process the document using enhanced processing
                    doc_result = self.process_document_for_storage(doc_file, {
                        'procedure_code': procedure_code,
                        'procedure_title': procedure_title
                    })

                    if doc_result and doc_result['content']:
                        # Update the source URL to point to the Supabase storage URL if upload was successful
                        storage_source_url = file_storage_url or actual_source_url or doc_file

                        # Add the file storage URL to the metadata
                        if not doc_result['metadata']:
                            doc_result['metadata'] = {}
                        if file_storage_url:
                            doc_result['metadata']['file_storage_url'] = file_storage_url
                            doc_result['metadata']['original_file_path'] = doc_file

                        # Store the full document content in knowledge base (not chunked)
                        success = vector_store.store_full_document_content(
                            ministry_name=ministry_name,
                            procedure_code=procedure_code,
                            procedure_title=procedure_title,
                            source_url=storage_source_url,
                            full_content=doc_result['content'],  # Store the full content as requested
                            metadata=doc_result['metadata']
                        )

                        if success:
                            processed_count += 1
                            if file_storage_url:
                                uploaded_count += 1
                                print(f"   ✅ Stored & uploaded: {procedure_code}")
                            else:
                                print(f"   ✅ Stored (no upload): {procedure_code}")
                        else:
                            failed_count += 1
                            print(f"   ❌ Failed to store: {procedure_code}")
                    else:
                        failed_count += 1
                        print(f"   ❌ Failed to process: {procedure_code}")

                except Exception as e:
                    failed_count += 1
                    print(f"   ❌ Error processing {doc_file}: {e}")
                    import traceback
                    traceback.print_exc()
                    continue

            print(f"   📊 Summary: {processed_count} processed, {uploaded_count} uploaded to storage, {failed_count} failed")
            return processed_count > 0

        except Exception as e:
            print(f"   ❌ Error in store_documents_with_embeddings: {e}")
            import traceback
            traceback.print_exc()
            return False

    def _add_link_to_word_file(self, filepath, url):
        """
        Add procedure link to Word file with proper error handling and backup mechanism
        """
        if not filepath or not os.path.exists(filepath):
            print(f"   ⚠️ File không tồn tại: {filepath}")
            return filepath  # Return original path, don't fail

        # Chỉ xử lý file .docx
        if not filepath.endswith('.docx'):
            print(f"   ⚠️ Chỉ hỗ trợ .docx, bỏ qua: {filepath}")
            return filepath

        # Tạo backup trước khi modify
        backup_path = filepath + '.backup'
        try:
            import shutil
            shutil.copy2(filepath, backup_path)
        except Exception as e:
            print(f"   ⚠️ Không thể tạo backup: {e}")
            return filepath

        try:
            # Validate file Word hợp lệ
            try:
                doc = Document(filepath)
            except Exception as e:
                print(f"   ⚠️ File Word không hợp lệ: {e}")
                return filepath

            # Thêm link vào cuối document
            doc.add_paragraph()  # Dòng trống

            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

            # Text trước link
            run1 = p.add_run("Để xem chi tiết thủ tục hành chính, vui lòng truy cập: ")
            run1.font.size = Pt(12)
            run1.font.name = 'Times New Roman'

            # Hyperlink
            run2 = p.add_run(url)
            run2.font.size = Pt(12)
            run2.font.name = 'Times New Roman'
            run2.font.color.rgb = RGBColor(0, 0, 255)
            run2.font.underline = True

            # Save document
            doc.save(filepath)

            # Xóa backup nếu thành công
            if os.path.exists(backup_path):
                os.remove(backup_path)

            print(f"   ✅ Đã thêm link vào: {os.path.basename(filepath)}")
            return filepath

        except PermissionError:
            print(f"   ⚠️ File đang được sử dụng, bỏ qua thêm link")
            # Restore backup
            if os.path.exists(backup_path):
                try:
                    shutil.move(backup_path, filepath)
                except:
                    pass
            return filepath

        except Exception as e:
            print(f"   ❌ Lỗi thêm link: {e}")
            # Restore backup
            if os.path.exists(backup_path):
                try:
                    shutil.move(backup_path, filepath)
                except:
                    pass
            return filepath

        finally:
            # Cleanup backup nếu còn tồn tại
            if os.path.exists(backup_path):
                try:
                    os.remove(backup_path)
                except:
                    pass

    def _create_supabase_storage_client(self):
        """
        Create a Supabase storage client for file uploads
        """
        try:
            from supabase import create_client
            import os

            supabase_url = os.getenv("SUPABASE_URL")
            supabase_key = os.getenv("SUPABASE_KEY")

            if not supabase_url or not supabase_key:
                print("   ❌ SUPABASE_URL and SUPABASE_KEY environment variables are required for file storage")
                return None

            client = create_client(supabase_url, supabase_key)
            return client.storage
        except ImportError:
            print("   ❌ Supabase library not installed. Run: pip install supabase")
            return None
        except Exception as e:
            print(f"   ❌ Error creating Supabase storage client: {e}")
            return None

    def upload_file_to_supabase_storage(self, file_path, bucket_name="documents", folder_path=None):
        """
        Upload a file to Supabase storage bucket

        Args:
            file_path (str): Path to the local file to upload
            bucket_name (str): Name of the Supabase storage bucket
            folder_path (str): Optional folder path within the bucket

        Returns:
            str or None: Public URL of the uploaded file, or None if upload failed
        """
        try:
            storage_client = self._create_supabase_storage_client()
            if not storage_client:
                return None

            # Prepare the file for upload
            file_name = os.path.basename(file_path)

            # If folder_path is provided, include it in the file path
            if folder_path:
                file_key = f"{folder_path.strip('/')}/{file_name}"
            else:
                file_key = file_name

            # Read the file content
            with open(file_path, 'rb') as f:
                file_content = f.read()

            # Set file options based on the file type
            file_extension = os.path.splitext(file_path)[1].lower()
            if file_extension == '.docx':
                content_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            elif file_extension == '.doc':
                content_type = 'application/msword'
            elif file_extension == '.pdf':
                content_type = 'application/pdf'
            else:
                content_type = 'application/octet-stream'  # Default

            file_options = {'content-type': content_type}

            # First, let's check if the requested bucket exists by attempting to list buckets
            bucket_exists = False
            try:
                # Test if we can access the bucket by trying to list it
                buckets_response = storage_client.list_buckets()
                if hasattr(buckets_response, 'data') and buckets_response.data:
                    bucket_exists = any(bucket.name == bucket_name for bucket in buckets_response.data)
                    print(f"   ℹ️  Found {len(buckets_response.data)} buckets in project")

                    # List all available buckets
                    available_buckets = [bucket.name for bucket in buckets_response.data]
                    print(f"   ℹ️  Available buckets: {available_buckets}")

                    if not bucket_exists:
                        print(f"   ⚠️  Bucket '{bucket_name}' does not exist.")
                        # Don't automatically fallback to 'public' since it might not exist either
                        # Instead, check if 'public' exists
                        if 'public' in available_buckets:
                            print(f"   ℹ️  'public' bucket exists, using it instead.")
                            bucket_name = 'public'
                            file_key = f"gov_procedures/{file_key}"
                        elif available_buckets:
                            # Use the first available bucket
                            fallback_bucket = available_buckets[0]
                            print(f"   ℹ️  Using first available bucket: '{fallback_bucket}'")
                            bucket_name = fallback_bucket
                            file_key = f"gov_procedures/{file_key}"
                        else:
                            print(f"   ❌ No buckets available in project. Please create a bucket in Supabase dashboard.")
                            return None
                else:
                    print(f"   ⚠️  No buckets found in project. Using 'public' bucket as default.")
                    bucket_name = 'public'  # Default fallback
                    file_key = f"gov_procedures/{file_key}"
            except Exception as bucket_check_error:
                print(f"   ⚠️  Could not check available buckets: {bucket_check_error}")
                print(f"   ℹ️  Proceeding with bucket '{bucket_name}', but it might not exist.")
                # Continue with original bucket name but warn user

            # Upload the file to Supabase storage
            response = storage_client.from_(bucket_name).upload(
                file_key,
                file_content,
                file_options
            )

            # Check if upload was successful
            if hasattr(response, 'error') and response.error:
                # If upload to 'public' bucket fails, it might be because it's not configured for this
                if bucket_name == 'public':
                    print(f"   ⚠️ Upload to 'public' bucket failed: {response.error}")
                    print(f"   ℹ️  Please create a dedicated bucket 'government-documents' in Supabase dashboard")
                    return None
                else:
                    print(f"   ❌ Upload failed: {response.error}")
                    return None

            # Generate public URL for the uploaded file
            public_url = storage_client.from_(bucket_name).get_public_url(file_key)

            print(f"   ✅ File uploaded to Supabase storage: {file_name}")
            print(f"   📄 Public URL: {public_url}")

            return public_url

        except Exception as e:
            print(f"   ❌ Error uploading file to Supabase storage: {e}")
            import traceback
            traceback.print_exc()
            return None

    def store_documents_with_embeddings_and_file_storage(self, ministry_name: str, source_url: str = None):
        """
        Process and store all downloaded documents to vector storage with embeddings and file storage
        This method uploads the actual files to Supabase storage and stores content in knowledge base
        """
        if VectorStorage is None:
            print("   ❌ VectorStorage module not available. Please install required dependencies.")
            return False

        # Check if environment variables are set for Supabase
        import os
        supabase_url = os.getenv("SUPABASE_URL")
        supabase_key = os.getenv("SUPABASE_KEY")

        if not supabase_url or not supabase_key:
            print("   ❌ SUPABASE_URL and SUPABASE_KEY environment variables are required")
            return False

        try:
            # Initialize vector storage
            vector_store = VectorStorage(supabase_url=supabase_url, supabase_key=supabase_key)

            # Find all docx files in the ministry folder
            safe_ministry_name = self._sanitize_filename(ministry_name)
            ministry_folder = os.path.join(self.download_dir, safe_ministry_name)
            if not os.path.exists(ministry_folder):
                print(f"   ❌ Ministry folder does not exist: {ministry_folder}")
                return False

            doc_files = []
            for root, dirs, files in os.walk(ministry_folder):
                for file in files:
                    if file.lower().endswith(('.doc', '.docx')):
                        doc_files.append(os.path.join(root, file))

            if not doc_files:
                print(f"   ❌ No .doc or .docx files found in {ministry_folder}")
                return False

            print(f"   📁 Found {len(doc_files)} .doc/.docx files to process")

            # Process and store each document
            processed_count = 0
            failed_count = 0
            uploaded_count = 0

            for doc_file in doc_files:
                try:
                    # Extract procedure info from the file path or name
                    filename = os.path.basename(doc_file)
                    procedure_code = os.path.splitext(filename)[0]  # Get name without extension
                    procedure_title = f"Thủ tục {procedure_code}"

                    # Get the procedure URL from the list file if available
                    # Use the properly sanitized ministry name
                    safe_ministry_name = self._sanitize_filename(ministry_name)
                    list_file = os.path.join(ministry_folder, f'danh_sach_{safe_ministry_name}.txt')
                    actual_source_url = source_url
                    if os.path.exists(list_file):
                        with open(list_file, 'r', encoding='utf-8') as f:
                            content = f.read()
                            # Try to find the URL for this specific procedure in the list file
                            import re
                            # Look for the procedure code in the list file and extract the URL
                            pattern = rf'\[{re.escape(procedure_code)}\].*?\n\s*URL:\s*(https?://[^\n]+)'
                            match = re.search(pattern, content, re.DOTALL | re.IGNORECASE)
                            if match:
                                actual_source_url = match.group(1).strip()

                    # Upload the file to Supabase storage
                    file_storage_url = self.upload_file_to_supabase_storage(
                        file_path=doc_file,
                        bucket_name="government-documents",  # Using a dedicated bucket for gov documents
                        folder_path=f"{safe_ministry_name}/procedures"
                    )

                    # Process the document using enhanced processing
                    doc_result = self.process_document_for_storage(doc_file, {
                        'procedure_code': procedure_code,
                        'procedure_title': procedure_title
                    })

                    if doc_result and doc_result['content']:
                        # Update the source URL to point to the Supabase storage URL if upload was successful
                        storage_source_url = file_storage_url or actual_source_url or doc_file

                        # Add the file storage URL to the metadata
                        if not doc_result['metadata']:
                            doc_result['metadata'] = {}
                        if file_storage_url:
                            doc_result['metadata']['file_storage_url'] = file_storage_url
                            doc_result['metadata']['original_file_path'] = doc_file

                        # Store the full document content in knowledge base (not chunked)
                        success = vector_store.store_full_document_content(
                            ministry_name=ministry_name,
                            procedure_code=procedure_code,
                            procedure_title=procedure_title,
                            source_url=storage_source_url,
                            full_content=doc_result['content'],  # Store the full content as requested
                            metadata=doc_result['metadata']
                        )

                        if success:
                            processed_count += 1
                            if file_storage_url:
                                uploaded_count += 1
                                print(f"   ✅ Stored & uploaded: {procedure_code}")
                            else:
                                print(f"   ✅ Stored (no upload): {procedure_code}")
                        else:
                            failed_count += 1
                            print(f"   ❌ Failed to store: {procedure_code}")
                    else:
                        failed_count += 1
                        print(f"   ❌ Failed to process: {procedure_code}")

                except Exception as e:
                    failed_count += 1
                    print(f"   ❌ Error processing {doc_file}: {e}")
                    import traceback
                    traceback.print_exc()
                    continue

            print(f"   📊 Summary: {processed_count} processed, {uploaded_count} uploaded to storage, {failed_count} failed")
            return processed_count > 0

        except Exception as e:
            print(f"   ❌ Error in store_documents_with_embeddings_and_file_storage: {e}")
            import traceback
            traceback.print_exc()
            return False

    def _count_vietnamese_words(self, text):
        """
        Count words in Vietnamese text (space-separated)
        """
        if not text:
            return 0
        # Vietnamese words are space-separated
        return len(text.split())

    def _download_file(self, url, filepath, max_retries=5, retry_delay=3):
        """Download file from URL with retry and validation"""
        # Ensure directory exists before downloading
        os.makedirs(os.path.dirname(filepath), exist_ok=True)

        for attempt in range(max_retries + 1):
            try:
                if not url.startswith('http'):
                    url = urljoin('https://thutuc.dichvucong.gov.vn', url)

                # Try downloading with timeout and appropriate headers
                headers = {
                    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36',
                    'Accept': 'application/msword,application/vnd.openxmlformats-officedocument.wordprocessingml.document,*/*',
                    'Accept-Language': 'vi,en-US;q=0.9,en;q=0.8',
                    'Referer': 'https://thutuc.dichvucong.gov.vn/',
                    'Connection': 'keep-alive'
                }
                response = self.session.get(url, headers=headers, timeout=45, stream=True, allow_redirects=True)

                # Check status code
                if response.status_code not in [200, 206]:  # 206 is partial content
                    if response.status_code in [502, 503, 504, 505]:  # Server issues, wait longer
                        delay = retry_delay * (attempt + 1)  # Exponential backoff
                    else:
                        delay = retry_delay

                    if attempt < max_retries:
                        print(f"   Warning - Attempt {attempt + 1}/{max_retries + 1}: Download failed (status {response.status_code}), retrying in {delay}s...")
                        time.sleep(delay)
                        continue
                    return False

                # Check content-type (should be file, not HTML error page)
                content_type = response.headers.get('content-type', '').lower()
                if 'text/html' in content_type and 'application' not in content_type:
                    if attempt < max_retries:
                        print(f"   Warning - Attempt {attempt + 1}/{max_retries + 1}: Got HTML instead of file, retrying in {retry_delay}s...")
                        time.sleep(retry_delay)
                        continue
                    return False

                # Download file with proper error handling
                try:
                    # Ensure the directory exists before writing
                    os.makedirs(os.path.dirname(filepath), exist_ok=True)
                    with open(filepath, 'wb') as f:
                        for chunk in response.iter_content(chunk_size=8192):
                            if chunk:
                                f.write(chunk)
                except OSError as e:
                    # Handle file system errors
                    if attempt < max_retries:
                        print(f"   Warning - Attempt {attempt + 1}/{max_retries + 1}: File system error: {str(e)}, retrying in {retry_delay}s...")
                        time.sleep(retry_delay)
                        continue
                    return False

                # Validate file
                if not os.path.exists(filepath):
                    if attempt < max_retries:
                        print(f"   Warning - Attempt {attempt + 1}/{max_retries + 1}: File doesn't exist after download, retrying...")
                        time.sleep(retry_delay)
                        continue
                    return False

                file_size = os.path.getsize(filepath)

                # File too small (< 500 bytes) might be an error
                if file_size < 500:
                    if os.path.exists(filepath):
                        os.remove(filepath)
                    if attempt < max_retries:
                        print(f"   Warning - Attempt {attempt + 1}/{max_retries + 1}: File too small ({file_size} bytes), retrying in {retry_delay}s...")
                        time.sleep(retry_delay)
                        continue
                    return False

                # Check if it's actually a file (read some header bytes)
                try:
                    with open(filepath, 'rb') as f:
                        header = f.read(8)
                        if len(header) < 4:
                            if os.path.exists(filepath):
                                os.remove(filepath)
                            if attempt < max_retries:
                                print(f"   Warning - Attempt {attempt + 1}/{max_retries + 1}: Invalid header, retrying in {retry_delay}s...")
                                time.sleep(retry_delay)
                                continue
                            return False
                except OSError as e:
                    if os.path.exists(filepath):
                        os.remove(filepath)
                    if attempt < max_retries:
                        print(f"   Warning - Attempt {attempt + 1}/{max_retries + 1}: Error reading file header: {str(e)[:50]}, retrying in {retry_delay}s...")
                        time.sleep(retry_delay)
                        continue
                    return False

                # If successful, exit the retry loop
                print(f"   OK - Download successful after {attempt + 1} attempts")
                return True

            except requests.exceptions.Timeout:
                if os.path.exists(filepath):
                    try:
                        os.remove(filepath)
                    except:
                        pass
                if attempt < max_retries:
                    print(f"   Warning - Attempt {attempt + 1}/{max_retries + 1}: Timeout when downloading file, retrying in {retry_delay}s...")
                    time.sleep(retry_delay)
                else:
                    print(f"   Failed - Timeout after {max_retries + 1} attempts")
            except requests.exceptions.RequestException as e:
                if os.path.exists(filepath):
                    try:
                        os.remove(filepath)
                    except:
                        pass
                if attempt < max_retries:
                    print(f"   Warning - Attempt {attempt + 1}/{max_retries + 1}: Request error: {str(e)[:50]}, retrying in {retry_delay}s...")
                    time.sleep(retry_delay)
                else:
                    print(f"   Failed - Request failed after {max_retries + 1} attempts: {str(e)[:100]}")
            except OSError as e:
                if os.path.exists(filepath):
                    try:
                        os.remove(filepath)
                    except:
                        pass
                if attempt < max_retries:
                    print(f"   Warning - Attempt {attempt + 1}/{max_retries + 1}: OS error: {str(e)[:50]}, retrying in {retry_delay}s...")
                    time.sleep(retry_delay)
                else:
                    print(f"   Failed - OS failed after {max_retries + 1} attempts: {str(e)[:100]}")
            except Exception as e:
                # Remove corrupt file if exists
                if os.path.exists(filepath):
                    try:
                        os.remove(filepath)
                    except:
                        pass

                if attempt < max_retries:
                    print(f"   Warning - Attempt {attempt + 1}/{max_retries + 1}: Exception when downloading: {str(e)[:50]}, retrying in {retry_delay}s...")
                    time.sleep(retry_delay)
                else:
                    print(f"   Failed - Failed after {max_retries + 1} attempts: {str(e)[:100]}")

        return False

    def download_procedure_parallel(self, procedure, ministry_folder):
        """Enhanced download with error tracking"""
        proc_id = procedure['id']
        proc_code = procedure['code']
        proc_title = procedure['title']

        # Create safe filename by removing invalid characters for Windows
        safe_code = self._sanitize_filename(proc_code)
        huong_dan_folder = os.path.join(ministry_folder, "huong_dan")
        os.makedirs(huong_dan_folder, exist_ok=True)

        # Check cache and file existence
        cache_key = f"{ministry_folder}_{proc_id}"
        cached_entry = self.cache.get(cache_key, {})

        doc_filename = f"{safe_code}.doc"
        docx_filename = f"{safe_code}.docx"
        doc_path = os.path.join(huong_dan_folder, doc_filename)
        docx_path = os.path.join(huong_dan_folder, docx_filename)

        final_filepath = None
        needs_download = True

        # Check if docx exists and is valid
        if os.path.exists(docx_path):
            current_size = os.path.getsize(docx_path)
            if current_size > 0 and (cached_entry.get('size') is None or current_size == cached_entry.get('size')):
                if cached_entry.get('checksum'):
                    current_checksum = self._calculate_file_checksum(docx_path)
                    if current_checksum and current_checksum == cached_entry.get('checksum'):
                        final_filepath = docx_path
                        needs_download = False
                        with self.stats_lock:
                            self.stats['cached'] += 1
                        print(f"   File already exists and validated: {proc_code}.docx (cached)")
                    else:
                        os.remove(docx_path)
                        print(f"   Checksum mismatch for {proc_code}.docx, re-downloading.")
                else:
                    final_filepath = docx_path
                    needs_download = False
                    with self.stats_lock:
                        self.stats['cached'] += 1
                    print(f"   File already exists: {proc_code}.docx (cached, no checksum validation)")
            else:
                if os.path.exists(docx_path): os.remove(docx_path)
                print(f"   File size mismatch for {proc_code}.docx, re-downloading.")

        # If docx doesn't exist or is invalid, check doc
        elif os.path.exists(doc_path):
            current_size = os.path.getsize(doc_path)
            if current_size > 0 and (cached_entry.get('size') is None or current_size == cached_entry.get('size')):
                if cached_entry.get('checksum'):
                    current_checksum = self._calculate_file_checksum(doc_path)
                    if current_checksum and current_checksum == cached_entry.get('checksum'):
                        final_filepath = doc_path
                        needs_download = False
                        with self.stats_lock:
                            self.stats['cached'] += 1
                        print(f"   File already exists and validated: {proc_code}.doc (cached)")
                    else:
                        os.remove(doc_path)
                        print(f"   Checksum mismatch for {proc_code}.doc, re-downloading.")
                else:
                    final_filepath = doc_path
                    needs_download = False
                    with self.stats_lock:
                        self.stats['cached'] += 1
                    print(f"   File already exists: {proc_code}.doc (cached, no checksum validation)")
            else:
                if os.path.exists(doc_path): os.remove(doc_path)
                print(f"   File size mismatch for {proc_code}.doc, re-downloading.")

        if needs_download:
            driver = None
            max_retries = 3  # FIX: Add retry mechanism for driver operations

            for retry in range(max_retries):
                try:
                    # Create driver for this thread
                    driver = self._create_driver()
                    wait = WebDriverWait(driver, 15)  # Increased timeout

                    # FIX: Navigate with retry and better error handling
                    try:
                        driver.get(procedure['url'])
                        # Wait for page to be fully loaded
                        wait.until(EC.presence_of_element_located((By.TAG_NAME, "body")))
                        time.sleep(1)  # Additional buffer for JS execution
                    except TimeoutException:
                        if retry < max_retries - 1:
                            print(f"   ⚠️ [{proc_code}] Page load timeout, retry {retry + 1}/{max_retries}")
                            if driver:
                                driver.quit()
                            time.sleep(2)
                            continue
                        else:
                            raise

                    # Download main guide file (.doc)
                    download_url = None

                    # FIX: More robust element finding with multiple strategies
                    for find_retry in range(3):
                        try:
                            # Strategy 1: Original XPath
                            try:
                                word_link = wait.until(EC.presence_of_element_located((
                                    By.XPATH,
                                    "//a[contains(@title, 'Tải xuống') and .//i[contains(@class, 'fa-file-word-o')]]"
                                )))
                                download_url = word_link.get_attribute('href')
                                if download_url:
                                    break
                            except TimeoutException:
                                pass

                            # Strategy 2: Find by icon class
                            try:
                                word_link = driver.find_element(
                                    By.XPATH,
                                    "//i[contains(@class, 'fa-file-word-o')]/ancestor::a"
                                )
                                download_url = word_link.get_attribute('href')
                                if download_url:
                                    break
                            except NoSuchElementException:
                                pass

                            # Strategy 3: Find any Word download link
                            try:
                                word_link = driver.find_element(
                                    By.CSS_SELECTOR,
                                    "a[href*='.doc'], a[href*='download']"
                                )
                                download_url = word_link.get_attribute('href')
                                if download_url and ('.doc' in download_url.lower()):
                                    break
                            except NoSuchElementException:
                                pass

                            if find_retry < 2:
                                time.sleep(1)

                        except Exception as e:
                            if find_retry < 2:
                                time.sleep(1)
                            else:
                                print(f"   ⚠️ [{proc_code}] Cannot find guide link after retries")

                    if download_url:
                        filename = f"{safe_code}.doc"
                        filepath = os.path.join(huong_dan_folder, filename)

                        # DEBUG: Print filepath to see what was constructed
                        print(f"DEBUG: Filepath for download: {filepath}")
                        print(f"DEBUG: Directory: {huong_dan_folder}")
                        print(f"DEBUG: Filename: {filename}")
                        print(f"DEBUG: Safe code: {safe_code}")

                        # Additional path validation
                        try:
                            # Ensure directory exists - this might also fail with invalid characters
                            os.makedirs(huong_dan_folder, exist_ok=True)
                        except OSError as e:
                            print(f"   ❌ Error creating directory {huong_dan_folder}: {e}")
                            with self.stats_lock:
                                self.stats['failed'] += 1
                            if driver:
                                driver.quit()
                            return False, proc_code, f"error: directory creation failed - {str(e)[:50]}"

                        # Use session to download
                        session = requests.Session()
                        session.headers.update({
                            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
                        })

                        # Download file
                        if self._download_file_with_session(download_url, filepath, session):
                            # First, convert .doc to .docx if necessary
                            converted_filepath = self._convert_doc_to_docx(filepath)

                            # Add link to the converted (or original) Word file
                            final_processed_filepath = self._add_link_to_word_file(converted_filepath, procedure['url'])

                            if final_processed_filepath:
                                # Calculate checksum for the final processed file
                                checksum = self._calculate_file_checksum(final_processed_filepath)
                                file_size = os.path.getsize(final_processed_filepath)
                                # Update cache
                                with self.cache_lock:
                                    self.cache[cache_key] = {
                                        'code': proc_code,
                                        'title': proc_title,
                                        'downloaded': True,
                                        'checksum': checksum,
                                        'size': file_size
                                    }
                                with self.stats_lock:
                                    self.stats['success'] += 1

                                # Check if actual conversion happened
                                if converted_filepath != filepath:
                                    print(f"   ✅ [{proc_code}] File downloaded, converted to .docx, and link added successfully")
                                else:
                                    print(f"   ✅ [{proc_code}] File downloaded (.doc format), link added successfully")
                                if driver:
                                    driver.quit()
                                return True, proc_code, "success"
                            else:
                                # If adding link fails, keep the file anyway
                                # Check if actual conversion happened
                                if converted_filepath != filepath:
                                    print(f"   ⚠️ [{proc_code}] File downloaded and converted to .docx, but couldn't add link (file still preserved)")
                                else:
                                    print(f"   ⚠️ [{proc_code}] File downloaded (.doc format), but couldn't add link (file still preserved)")
                                checksum = self._calculate_file_checksum(converted_filepath)  # Use converted filepath
                                file_size = os.path.getsize(converted_filepath)
                                # Update cache with file (converted or original)
                                with self.cache_lock:
                                    self.cache[cache_key] = {
                                        'code': proc_code,
                                        'title': proc_title,
                                        'downloaded': True,
                                        'checksum': checksum,
                                        'size': file_size
                                    }
                                with self.stats_lock:
                                    self.stats['success'] += 1  # Count as success since file was processed
                                if driver:
                                    driver.quit()
                                return True, proc_code, "success_download_only"  # Changed to reflect that link addition failed

                        if driver:
                            driver.quit()
                        break  # Success, exit retry loop
                    else:
                        if retry < max_retries - 1:
                            print(f"   ⚠️ [{proc_code}] No download link, retry {retry + 1}/{max_retries}")
                            if driver:
                                driver.quit()
                            time.sleep(2)
                            continue
                        else:
                            with self.stats_lock:
                                self.stats['failed'] += 1
                            return False, proc_code, "no_download_link"

                except TimeoutException as e:
                    if retry < max_retries - 1:
                        print(f"   ⚠️ [{proc_code}] Timeout error, retry {retry + 1}/{max_retries}")
                        if driver:
                            try:
                                driver.quit()
                            except:
                                pass
                        time.sleep(2)
                        continue
                    else:
                        with self.stats_lock:
                            self.stats['failed'] += 1
                        if driver:
                            try:
                                driver.quit()
                            except:
                                pass
                        # FIX: Track the error
                        error_msg = f"timeout_error: {str(e)[:50]}"
                        self.error_tracker.add_error('timeout', proc_code, error_msg)
                        self.logger.error(f"[{proc_code}] {error_msg}")
                        return False, proc_code, f"timeout_error: {str(e)[:50]}"

                except Exception as e:
                    if retry < max_retries - 1:
                        print(f"   ⚠️ [{proc_code}] Error: {str(e)[:50]}, retry {retry + 1}/{max_retries}")
                        if driver:
                            try:
                                driver.quit()
                            except:
                                pass
                        time.sleep(2)
                        continue
                    else:
                        with self.stats_lock:
                            self.stats['failed'] += 1
                        if driver:
                            try:
                                driver.quit()
                            except:
                                pass
                        # FIX: Track the error
                        error_msg = f"general_error: {str(e)[:50]}"
                        # Categorize error
                        if 'path' in error_msg.lower() or 'PATH' in str(e):
                            self.error_tracker.add_error('path_length', proc_code, error_msg)
                        elif 'write-protected' in error_msg.lower() or 'disk' in error_msg.lower():
                            self.error_tracker.add_error('permission', proc_code, error_msg)
                        elif 'conversion' in error_msg.lower():
                            self.error_tracker.add_error('conversion', proc_code, error_msg)
                        else:
                            self.error_tracker.add_error('other', proc_code, error_msg)
                        self.logger.error(f"[{proc_code}] {error_msg}")
                        return False, proc_code, f"error: {str(e)[:50]}"

            # FIX: Always cleanup driver
            if driver:
                try:
                    driver.quit()
                except:
                    pass
        else:
            # If not downloading (i.e., cached), ensure it's converted to docx and has link
            current_filepath = final_filepath
            # Convert if needed (this may change .doc to .docx)
            converted_filepath = self._convert_doc_to_docx(current_filepath)

            # The converted filepath may be different (e.g., .doc converted to .docx)
            # Update final_filepath to the actual file that exists
            if os.path.exists(converted_filepath):
                final_filepath = converted_filepath
            else:
                # If conversion failed, stick with the original file
                final_filepath = current_filepath

            # Add link to both .doc and .docx files if the add link function supports it
            if final_filepath and (final_filepath.endswith('.docx') or final_filepath.endswith('.doc')):
                # Report the actual file format after potential conversion
                if converted_filepath != current_filepath and final_filepath.endswith('.docx'):
                    print(f"   Adding link to cached .docx file for {proc_code} (was previously .doc)...")
                elif current_filepath.endswith('.docx'):
                    print(f"   Adding link to cached .docx file for {proc_code}...")
                else:
                    print(f"   Adding link to cached .doc file for {proc_code}...")

                final_processed_filepath = self._add_link_to_word_file(final_filepath, procedure['url'])

                if final_processed_filepath:
                    # Update checksum/size in cache after potential link addition
                    with self.cache_lock:
                        if cache_key in self.cache:  # Check if cache entry exists
                            self.cache[cache_key]['checksum'] = self._calculate_file_checksum(final_processed_filepath)
                            self.cache[cache_key]['size'] = os.path.getsize(final_processed_filepath)
                        else:
                            # If cache entry doesn't exist, create one
                            self.cache[cache_key] = {
                                'code': proc_code,
                                'title': proc_title,
                                'downloaded': True,
                                'checksum': self._calculate_file_checksum(final_processed_filepath),
                                'size': os.path.getsize(final_processed_filepath)
                            }
                else:
                    print(f"   Warning: Could not add link to cached file {proc_code}.")
            else:
                print(f"   Warning: Cached file {proc_code} not a Word file, skipping link add.")
            return True, proc_code, "cached"

    def _download_file_with_session(self, url, filepath, session, max_retries=5, retry_delay=3):
        """Download file from URL with specific session, with retry on error"""
        # Ensure directory exists before downloading
        os.makedirs(os.path.dirname(filepath), exist_ok=True)

        for attempt in range(max_retries + 1):
            try:
                if not url.startswith('http'):
                    url = urljoin('https://thutuc.dichvucong.gov.vn', url)

                # Try downloading with timeout and appropriate headers
                headers = {
                    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36',
                    'Accept': 'application/msword,application/vnd.openxmlformats-officedocument.wordprocessingml.document,*/*',
                    'Accept-Language': 'vi,en-US;q=0.9,en;q=0.8',
                    'Referer': 'https://thutuc.dichvucong.gov.vn/',
                    'Connection': 'keep-alive'
                }
                response = session.get(url, headers=headers, timeout=45, stream=True, allow_redirects=True)

                # Check status code
                if response.status_code not in [200, 206]:  # 206 is partial content
                    if response.status_code in [502, 503, 504, 505]:  # Server issues, wait longer
                        delay = retry_delay * (attempt + 1)  # Exponential backoff
                    else:
                        delay = retry_delay

                    if attempt < max_retries:
                        print(f"   Warning - Attempt {attempt + 1}/{max_retries + 1}: Download failed (status {response.status_code}), retrying in {delay}s...")
                        time.sleep(delay)
                        continue
                    return False

                # Check content-type (should be file, not HTML error page)
                content_type = response.headers.get('content-type', '').lower()
                if 'text/html' in content_type and 'application' not in content_type:
                    if attempt < max_retries:
                        print(f"   Warning - Attempt {attempt + 1}/{max_retries + 1}: Got HTML instead of file, retrying in {retry_delay}s...")
                        time.sleep(retry_delay)
                        continue
                    return False

                # Download file with proper error handling
                try:
                    # Ensure the directory exists before writing
                    os.makedirs(os.path.dirname(filepath), exist_ok=True)
                    with open(filepath, 'wb') as f:
                        for chunk in response.iter_content(chunk_size=8192):
                            if chunk:
                                f.write(chunk)
                except OSError as e:
                    # Handle file system errors
                    if attempt < max_retries:
                        print(f"   Warning - Attempt {attempt + 1}/{max_retries + 1}: File system error: {str(e)}, retrying in {retry_delay}s...")
                        time.sleep(retry_delay)
                        continue
                    return False

                # Validate file
                if not os.path.exists(filepath):
                    if attempt < max_retries:
                        print(f"   Warning - Attempt {attempt + 1}/{max_retries + 1}: File doesn't exist after download, retrying...")
                        time.sleep(retry_delay)
                        continue
                    return False

                file_size = os.path.getsize(filepath)

                # File too small (< 500 bytes) might be an error
                if file_size < 500:
                    if os.path.exists(filepath):
                        os.remove(filepath)
                    if attempt < max_retries:
                        print(f"   Warning - Attempt {attempt + 1}/{max_retries + 1}: File too small ({file_size} bytes), retrying in {retry_delay}s...")
                        time.sleep(retry_delay)
                        continue
                    return False

                # Check if it's actually a file (read some header bytes)
                try:
                    with open(filepath, 'rb') as f:
                        header = f.read(8)
                        if len(header) < 4:
                            if os.path.exists(filepath):
                                os.remove(filepath)
                            if attempt < max_retries:
                                print(f"   Warning - Attempt {attempt + 1}/{max_retries + 1}: Invalid header, retrying in {retry_delay}s...")
                                time.sleep(retry_delay)
                                continue
                            return False
                except OSError as e:
                    if os.path.exists(filepath):
                        os.remove(filepath)
                    if attempt < max_retries:
                        print(f"   Warning - Attempt {attempt + 1}/{max_retries + 1}: Error reading file header: {str(e)[:50]}, retrying in {retry_delay}s...")
                        time.sleep(retry_delay)
                        continue
                    return False

                # If successful, exit the retry loop
                print(f"   OK - Download successful after {attempt + 1} attempts")
                return True

            except requests.exceptions.Timeout:
                if os.path.exists(filepath):
                    try:
                        os.remove(filepath)
                    except:
                        pass
                if attempt < max_retries:
                    print(f"   Warning - Attempt {attempt + 1}/{max_retries + 1}: Timeout when downloading file, retrying in {retry_delay}s...")
                    time.sleep(retry_delay)
                else:
                    print(f"   Failed - Timeout after {max_retries + 1} attempts")
            except requests.exceptions.RequestException as e:
                if os.path.exists(filepath):
                    try:
                        os.remove(filepath)
                    except:
                        pass
                if attempt < max_retries:
                    print(f"   Warning - Attempt {attempt + 1}/{max_retries + 1}: Request error: {str(e)[:50]}, retrying in {retry_delay}s...")
                    time.sleep(retry_delay)
                else:
                    print(f"   Failed - Request failed after {max_retries + 1} attempts: {str(e)[:100]}")
            except OSError as e:
                if os.path.exists(filepath):
                    try:
                        os.remove(filepath)
                    except:
                        pass
                if attempt < max_retries:
                    print(f"   Warning - Attempt {attempt + 1}/{max_retries + 1}: OS error: {str(e)[:50]}, retrying in {retry_delay}s...")
                    time.sleep(retry_delay)
                else:
                    print(f"   Failed - OS failed after {max_retries + 1} attempts: {str(e)[:100]}")
            except Exception as e:
                # Remove corrupt file if exists
                if os.path.exists(filepath):
                    try:
                        os.remove(filepath)
                    except:
                        pass

                if attempt < max_retries:
                    print(f"   Warning - Attempt {attempt + 1}/{max_retries + 1}: Exception when downloading: {str(e)[:50]}, retrying in {retry_delay}s...")
                    time.sleep(retry_delay)
                else:
                    print(f"   Failed - Failed after {max_retries + 1} attempts: {str(e)[:100]}")

        return False

    def download_all_guides(self, procedures, ministry_folder):
        """Download all guide files with multi-threading"""
        print("="*70)
        print("⚡ BẮT ĐẦU DOWNLOAD FILE HƯỚNG DẪN (ĐA LUỒNG)")
        print("="*70)

        # First pass: download all procedures
        with ThreadPoolExecutor(max_workers=self.max_workers) as executor:
            # Submit all tasks
            future_to_procedure = {
                executor.submit(self.download_procedure_parallel, procedure, ministry_folder): procedure
                for procedure in procedures
            }

            # Track progress
            completed = 0
            total = len(procedures)

            for future in as_completed(future_to_procedure):
                completed += 1
                success, proc_code, status = future.result()

                if status == "cached":
                    print(f"\n📝 [{completed}/{total}] [{proc_code}] ĐÃ CÓ TRONG CACHE")
                elif success:
                    print(f"\n✅ [{completed}/{total}] [{proc_code}] Download thành công")
                else:
                    print(f"\n❌ [{completed}/{total}] [{proc_code}] Download thất bại ({status})")

                # Print stats every 5 procedures or at the end
                if completed % 5 == 0 or completed == total:
                    print(f"📊 Tiến độ: {completed}/{total} | "
                          f"✅ {self.stats['success']} | "
                          f"💾 {self.stats['cached']} | "
                          f"❌ {self.stats['failed']}")

        # Summary of first run
        print(f"\n📊 KẾT QUẢ SAU LẦN CHẠY ĐẦU:")
        print(f"   ✅ Thành công: {self.stats['success']}")
        print(f"   💾 Đã có (cache): {self.stats['cached']}")
        print(f"   ❌ Thất bại: {self.stats['failed']}")

        if self.stats['failed'] > 0:
            print(f"ℹ️  Failed downloads have built-in retry mechanisms")

        print(f"\n🎉 HOÀN TẤT TIẾN TRÌNH TẢI XUỐNG!")

    def export_chunks_to_csv(self, ministry_name):
        """Export processed full content to a CSV file for verification"""
        import csv
        import os
        from pathlib import Path

        safe_ministry_name = self._sanitize_filename(ministry_name)
        ministry_folder = os.path.join(self.download_dir, safe_ministry_name)

        if not os.path.exists(ministry_folder):
            print(f"   ❌ Ministry folder does not exist: {ministry_folder}")
            return False

        # Find all doc/docx files in the ministry folder
        doc_files = []
        for root, dirs, files in os.walk(ministry_folder):
            for file in files:
                if file.lower().endswith(('.doc', '.docx')):
                    doc_files.append(os.path.join(root, file))

        if not doc_files:
            print(f"   ❌ No .doc or .docx files found in {ministry_folder}")
            return False

        print(f"   📁 Found {len(doc_files)} .doc/.docx files to process for CSV export")

        # Prepare CSV file
        csv_filename = os.path.join(self.download_dir, f"{safe_ministry_name}_documents_verification.csv")

        with open(csv_filename, 'w', newline='', encoding='utf-8') as csvfile:
            fieldnames = ['procedure_code', 'procedure_title', 'original_filename', 'content_length', 'full_content_preview', 'metadata']
            writer = csv.DictWriter(csvfile, fieldnames=fieldnames)

            writer.writeheader()

            # Process each document to extract and export full content
            processed_count = 0
            failed_count = 0

            for doc_file in doc_files:
                try:
                    # Extract procedure info from the file path or name
                    filename = os.path.basename(doc_file)
                    procedure_code = os.path.splitext(filename)[0]  # Get name without extension
                    procedure_title = f"Thủ tục {procedure_code}"

                    # Process the document to get full content using the existing method
                    doc_result = self.process_document_for_storage(doc_file, {
                        'procedure_code': procedure_code,
                        'procedure_title': procedure_title
                    })

                    if doc_result and doc_result['content']:
                        # Write full content to CSV (truncated for preview)
                        writer.writerow({
                            'procedure_code': procedure_code,
                            'procedure_title': procedure_title,
                            'original_filename': filename,
                            'content_length': len(doc_result['content']),
                            'full_content_preview': doc_result['content'][:500] + "..." if len(doc_result['content']) > 500 else doc_result['content'],  # Truncate for preview
                            'metadata': str(doc_result['metadata'])
                        })

                        processed_count += 1
                        print(f"   ✅ Exported full content from: {filename} (length: {len(doc_result['content'])} chars)")
                    else:
                        failed_count += 1
                        print(f"   ❌ Failed to process for export: {filename}")

                except Exception as e:
                    failed_count += 1
                    print(f"   ❌ Error processing {doc_file} for export: {e}")
                    import traceback
                    traceback.print_exc()
                    continue

        print(f"\n📊 Export Summary: {processed_count} documents processed, {failed_count} failed")
        print(f"📄 CSV file saved as: {csv_filename}")
        print(f"📈 This CSV contains full content for verification before storing to Supabase")
        return processed_count > 0

    def export_to_excel(self, ministry_name):
        """
        Export processed documents to Excel file with proper Unicode support
        """
        try:
            import pandas as pd
            import os
            from pathlib import Path
        except ImportError:
            print("   ❌ pandas library not installed. Run: pip install pandas openpyxl")
            return False

        safe_ministry_name = self._sanitize_filename(ministry_name)
        ministry_folder = os.path.join(self.download_dir, safe_ministry_name)

        if not os.path.exists(ministry_folder):
            print(f"   ❌ Ministry folder does not exist: {ministry_folder}")
            return False

        # Find all doc/docx files in the ministry folder
        doc_files = []
        for root, dirs, files in os.walk(ministry_folder):
            for file in files:
                if file.lower().endswith(('.doc', '.docx')):
                    doc_files.append(os.path.join(root, file))

        if not doc_files:
            print(f"   ❌ No .doc or .docx files found in {ministry_folder}")
            return False

        print(f"   📁 Found {len(doc_files)} .doc/.docx files to process for Excel export")

        # Prepare data list
        data_list = []
        processed_count = 0
        failed_count = 0

        for doc_file in doc_files:
            try:
                # Extract procedure info from the file path or name
                filename = os.path.basename(doc_file)
                procedure_code = os.path.splitext(filename)[0]  # Get name without extension
                procedure_title = f"Thủ tục {procedure_code}"

                # Process the document using enhanced processing
                doc_result = self.process_document_for_storage(doc_file, {
                    'procedure_code': procedure_code,
                    'procedure_title': procedure_title
                })

                if doc_result and doc_result['content']:
                    # Create a record for Excel
                    record = {
                        'procedure_code': procedure_code,
                        'procedure_title': procedure_title,
                        'original_filename': filename,
                        'content_length': len(doc_result['content']),
                        'full_content': doc_result['content'],  # Full content in Excel
                        'full_content_preview': doc_result['content'][:500] + "..." if len(doc_result['content']) > 500 else doc_result['content'],
                        'metadata': str(doc_result['metadata'])
                    }
                    data_list.append(record)
                    processed_count += 1
                    print(f"   ✅ Processed for Excel: {filename} (length: {len(doc_result['content'])} chars)")
                else:
                    failed_count += 1
                    print(f"   ❌ Failed to process for Excel export: {filename}")

            except Exception as e:
                failed_count += 1
                print(f"   ❌ Error processing {doc_file} for Excel export: {e}")
                import traceback
                traceback.print_exc()
                continue

        if data_list:
            # Create Excel file
            excel_filename = os.path.join(self.download_dir, f"{safe_ministry_name}_documents_verification.xlsx")

            # Create DataFrame and save to Excel
            df = pd.DataFrame(data_list)
            df.to_excel(excel_filename, index=False, engine='openpyxl')

            print(f"\n📊 Excel Export Summary: {processed_count} documents processed, {failed_count} failed")
            print(f"📚 Excel file saved as: {excel_filename}")
            print(f"📈 Excel contains full content with proper Unicode support")
            return processed_count > 0
        else:
            print(f"\n❌ No data to export to Excel")
            return False

    def run(self, url, ministry_name, max_procedures=None, store_to_vector=False):
        """Enhanced run with error reporting"""
        start_time = time.time()

        try:
            # Create ministry folder
            safe_ministry = self._sanitize_filename(ministry_name)
            ministry_folder = os.path.join(self.download_dir, safe_ministry)
            os.makedirs(ministry_folder, exist_ok=True)

            # Get procedure list
            procedures = self.get_all_procedures(url, ministry_name)

            if not procedures:
                print("❌ Không tìm thấy thủ tục nào!")
                return

            # Limit if needed
            if max_procedures:
                procedures = procedures[:max_procedures]
                print(f"⚠️ Giới hạn: {max_procedures} thủ tục đầu tiên\n")

            # Save list with full procedure names
            list_file = os.path.join(ministry_folder, f'danh_sach_{safe_ministry}.txt')
            with open(list_file, 'w', encoding='utf-8') as f:
                f.write(f"DANH SÁCH THỦ TỤC HÀNH CHÍNH - {ministry_name.upper()}\n")
                f.write("="*80 + "\n")
                f.write(f"Tổng số: {len(procedures)} thủ tục\n")
                f.write(f"Ngày tạo: {time.strftime('%Y-%m-%d %H:%M:%S')}\n")
                f.write("="*80 + "\n\n")
                for i, p in enumerate(procedures, 1):
                    f.write(f"{i}. [{p['code']}] {p['title']}\n")
                    f.write(f"   URL: {p['url']}\n\n")

            print(f"✅ Đã lưu danh sách: {list_file}")
            print(f"📁 Cấu trúc thư mục:")
            print(f"   {ministry_folder}/")
            print(f"   ├── huong_dan/     (Hướng dẫn .doc/.docx)")
            print(f"   └── danh_sach_{safe_ministry}.txt\n")

            # Download files
            self.download_all_guides(procedures, ministry_folder)

            # FIX: Save error report
            error_summary = self.error_tracker.get_summary()
            if error_summary:
                print("\n📊 ERROR SUMMARY:")
                for category, count in error_summary.items():
                    print(f"   {category}: {count} errors")

                # Save detailed report
                error_report_file = os.path.join(ministry_folder, 'error_report.txt')
                self.error_tracker.save_error_report(error_report_file)
                print(f"\n📄 Detailed error report saved: {error_report_file}")

            # Export chunks to CSV for verification
            print("\n📊 XUẤT CHUNKS RA FILE CSV ĐỂ KIỂM TRA...")
            self.export_chunks_to_csv(ministry_name)

            # Export to Excel for proper Unicode support and better formatting
            print("\n📊 XUẤT DỮ LIỆU RA FILE EXCEL (HỖ TRỢ TIẾNG VIỆT)...")
            self.export_to_excel(ministry_name)

            # Optionally store documents to vector storage with embeddings and file storage
            if store_to_vector and VectorStorage is not None:
                print("\n🔄 BẮT ĐẦU LƯU DOCUMENTS VÀO VECTOR STORAGE VÀ UPLOAD LÊN SUPABASE...")
                self.store_documents_with_embeddings_and_file_storage(ministry_name, url)

            # Import the generated CSV to Supabase table (this works even if store_to_vector failed due to missing buckets)
            print("\n🔄 BẮT ĐẦU IMPORT DỮ LIỆU CSV VÀO SUPABASE...")
            self.import_csv_to_supabase(ministry_name=ministry_name)

            # Results
            elapsed = time.time() - start_time

            print("\n" + "="*70)
            print("📊 KẾT QUẢ CUỐI CÙNG")
            print("="*70)
            print(f"✅ Thủ tục thành công: {self.stats['success']}")
            print(f"💾 Đã có (cache): {self.stats['cached']}")
            print(f"❌ Thất bại: {self.stats['failed']}")
            print(f"📁 Thư mục: {ministry_folder}/")
            print(f"⏱️ Thời gian: {elapsed/60:.1f} phút")

        except Exception as e:
            self.logger.error(f"Critical error: {e}", exc_info=True)
            print(f"\n❌ Lỗi nghiêm trọng: {e}")

        finally:
            try:
                self.driver.quit()
            except:
                pass
            print("\n🎉 HOÀN TẤT!")

    def export_all_documents_to_csv(self, ministry_name, csv_filename=None):
        """
        Extract content from all doc/docx files and export to a complete CSV file
        """
        import csv
        import os
        from pathlib import Path

        if csv_filename is None:
            safe_ministry_name = self._sanitize_filename(ministry_name)
            csv_filename = f"{safe_ministry_name}_documents.csv"

        # Find all doc/docx files in the ministry folder
        safe_ministry_name = self._sanitize_filename(ministry_name)
        ministry_folder = os.path.join(self.download_dir, safe_ministry_name)

        if not os.path.exists(ministry_folder):
            print(f"   ❌ Ministry folder does not exist: {ministry_folder}")
            return False

        doc_files = []
        for root, dirs, files in os.walk(ministry_folder):
            for file in files:
                if file.lower().endswith(('.doc', '.docx')):
                    doc_files.append(os.path.join(root, file))

        if not doc_files:
            print(f"   ❌ No .doc or .docx files found in {ministry_folder}")
            return False

        print(f"   Found {len(doc_files)} .doc/.docx files to process")

        # Prepare CSV file
        with open(csv_filename, 'w', newline='', encoding='utf-8') as csvfile:
            fieldnames = ['file_path', 'filename', 'procedure_code', 'procedure_title', 'content_length', 'content', 'metadata']
            writer = csv.DictWriter(csvfile, fieldnames=fieldnames)

            writer.writeheader()

            # Process each document
            processed_count = 0
            failed_count = 0

            for doc_file in doc_files:
                try:
                    # Extract procedure info from the file path or name
                    filename = os.path.basename(doc_file)
                    procedure_code = os.path.splitext(filename)[0]  # Get name without extension
                    procedure_title = f"Thủ tục {procedure_code}"

                    # Extract content from the document
                    content = self.extract_content_from_docx_path(doc_file)

                    # Extract metadata from the content
                    metadata = self.extract_structural_info_from_content(content)

                    # Write to CSV
                    writer.writerow({
                        'file_path': doc_file,
                        'filename': filename,
                        'procedure_code': procedure_code,
                        'procedure_title': procedure_title,
                        'content_length': len(content),
                        'content': content,
                        'metadata': str(metadata)  # Convert metadata dict to string for CSV
                    })

                    processed_count += 1
                    print(f"   Exported: {filename}")

                except Exception as e:
                    failed_count += 1
                    print(f"   Error processing {doc_file}: {e}")
                    import traceback
                    traceback.print_exc()
                    continue

        print(f"   Export Summary: {processed_count} processed, {failed_count} failed")
        print(f"   CSV file saved as: {csv_filename}")
        return processed_count > 0

    def import_csv_to_supabase(self, csv_filename=None, ministry_name=None):
        """
        Import CSV data directly to Supabase table
        Specifically designed to handle the documents verification CSV format with unique procedure records
        """
        import csv
        import os
        import json
        from pathlib import Path
        import hashlib

        # Check for required environment variables
        supabase_url = os.getenv("SUPABASE_URL")
        supabase_key = os.getenv("SUPABASE_KEY")

        if not supabase_url or not supabase_key:
            print("   ❌ SUPABASE_URL and SUPABASE_KEY environment variables are required for Supabase import")
            return False

        # If no CSV filename provided, try to find the verification CSV
        if not csv_filename:
            if ministry_name:
                safe_ministry_name = self._sanitize_filename(ministry_name)
                csv_filename = os.path.join(self.download_dir, f"{safe_ministry_name}_documents_verification.csv")
            else:
                # Try to find the documents verification CSV (not chunks) in the download directory
                import glob
                # Focus only on documents_verification.csv files (not chunks_verification.csv)
                verification_csvs = glob.glob(os.path.join(self.download_dir, "*_documents_verification.csv"))
                if verification_csvs:
                    csv_filename = verification_csvs[0]  # Use the first one
                else:
                    print("   ❌ No documents verification CSV file found. Run export_chunks_to_csv first.")
                    return False

        if not os.path.exists(csv_filename):
            print(f"   ❌ CSV file does not exist: {csv_filename}")
            return False

        try:
            # Import supabase client
            from supabase import create_client
            client = create_client(supabase_url, supabase_key)

            print(f"   ✅ Connected to Supabase, importing from: {csv_filename}")

            # Read the CSV file with proper encoding
            with open(csv_filename, 'r', encoding='utf-8-sig') as f:
                reader = csv.DictReader(f)
                rows = list(reader)

            print(f"   ℹ️  Found {len(rows)} rows to import to Supabase")

            # Process and import each row individually to avoid batch conflicts
            imported_count = 0
            failed_count = 0

            for row in rows:
                try:
                    # Parse metadata from string to JSON if needed
                    metadata_str = row.get('metadata', '{}')
                    if isinstance(metadata_str, str) and metadata_str.strip():
                        try:
                            # Use a more robust JSON parsing approach
                            import re
                            # Clean up the metadata string to handle common problematic patterns
                            clean_metadata_str = metadata_str.strip()
                            # Fix common JSON issues
                            clean_metadata_str = clean_metadata_str.replace('\n', '\\n').replace('\r', '\\r').replace('\t', '\\t')
                            metadata_json = json.loads(clean_metadata_str)
                        except json.JSONDecodeError:
                            # If JSON parsing fails, store it as a raw string in json object
                            metadata_json = {"raw": metadata_str}
                    else:
                        metadata_json = {}

                    # Convert content_length to integer
                    content_length_str = row.get('content_length', '0')
                    try:
                        content_length = int(content_length_str) if content_length_str.strip() else 0
                    except ValueError:
                        content_length = 0

                    # Get the full content (in verification CSV this is in full_content_preview)
                    full_content = row.get('full_content_preview', '') or row.get('content', '')

                    # Calculate document hash from content for duplicate detection
                    content_bytes = full_content.encode('utf-8', errors='ignore') if full_content else b''
                    doc_hash = hashlib.sha256(content_bytes).hexdigest() if content_bytes else ''

                    # Determine ministry name - if not provided, try to extract from CSV filename
                    if not ministry_name:
                        # Extract ministry name from the CSV filename
                        csv_basename = os.path.basename(csv_filename)
                        ministry_name_extracted = csv_basename.replace('_documents_verification.csv', '').replace('_', ' ')
                    else:
                        ministry_name_extracted = ministry_name

                    # Prepare the record for insertion
                    record = {
                        'procedure_code': row.get('procedure_code', ''),
                        'procedure_title': row.get('procedure_title', ''),
                        'original_filename': row.get('original_filename', ''),
                        'content_length': content_length,
                        'full_content': full_content,  # This contains the structured content from the CSV
                        'full_content_preview': full_content[:10000] if full_content else '',  # First 10k chars for preview
                        'metadata': metadata_json,
                        'ministry_name': ministry_name_extracted,
                        'doc_hash': doc_hash,
                        'file_size': content_length,
                        'source_url': f"file:///{row.get('original_filename', '')}"  # Add source URL for reference
                    }

                    # Try to insert the record individually to avoid batch conflicts
                    try:
                        response = client.table('government_procedures_knowledge').insert(record).execute()
                        imported_count += 1
                        if imported_count % 20 == 0:  # Print progress every 20 records
                            print(f"   ✅ {imported_count}/{len(rows)} records imported")
                    except Exception as insert_error:
                        error_msg = str(insert_error).lower()

                        # Check if it's a duplicate key error (doc_hash already exists)
                        if "duplicate key value violates unique constraint" in error_msg:
                            print(f"   ⚠️  Skipped duplicate procedure: {row.get('procedure_code', 'Unknown')}")
                            continue
                        elif "on conflict do update command cannot affect row a second time" in error_msg:
                            # This happens when batch inserting duplicates, shouldn't happen with individual inserts
                            print(f"   ⚠️  Duplicate conflict for procedure: {row.get('procedure_code', 'Unknown')}")
                            continue
                        elif "invalid input syntax for type json" in error_msg:
                            # Handle JSON parsing errors by storing raw metadata
                            print(f"   ⚠️  JSON parsing error for procedure {row.get('procedure_code', 'Unknown')}, using raw metadata")
                            record['metadata'] = {"raw": metadata_str}
                            # Retry insert with raw metadata
                            response = client.table('government_procedures_knowledge').insert(record).execute()
                            imported_count += 1
                        else:
                            print(f"   ❌ Error inserting record for procedure {row.get('procedure_code', 'Unknown')}: {insert_error}")
                            failed_count += 1
                            continue

                except Exception as e:
                    print(f"   ❌ Error processing row for procedure {row.get('procedure_code', 'Unknown')}: {e}")
                    import traceback
                    traceback.print_exc()
                    failed_count += 1
                    continue

            print(f"\n   🎉 Import completed: {imported_count} successful, {failed_count} failed")
            print(f"   📊 Data is now available in the government_procedures_knowledge table in Supabase")
            return imported_count > 0

        except ImportError:
            print("   ❌ Supabase library not installed. Run: pip install supabase")
            return False
        except Exception as e:
            print(f"   ❌ Error importing data to Supabase: {e}")
            import traceback
            traceback.print_exc()
            return False


def get_user_input(prompt, default_value=None, valid_options=None):
    """
    Get user input with fallback for environments where input() might not work
    """
    import sys
    import os

    # Check if stdin is a terminal (TTY)
    if not os.isatty(sys.stdin.fileno()):
        print(f"\n⚠️ Cannot get input from user (not a terminal environment). Using default value: {default_value}")
        return default_value

    try:
        choice = input(prompt).strip()
        if not choice and default_value is not None:
            return default_value
        if valid_options and choice not in valid_options:
            print(f"Choice '{choice}' is invalid. Using default value: {default_value}")
            return default_value
        return choice
    except (EOFError, KeyboardInterrupt):
        # Handle case where no input is available (like when running in certain IDEs)
        print(f"\n⚠️ Cannot get user input. Using default value: {default_value}")
        return default_value
    except Exception:
        # Fallback for other environments where input() doesn't work
        print(f"\n⚠️ Cannot get user input. Using default value: {default_value}")
        return default_value


def main():
    """Main function"""
    print("""
╔═══════════════════════════════════════════════════════════════╗
║   ENHANCED SCRAPER THỦ TỤC HÀNH CHÍNH - TẤT CẢ BỘ/NGÀNH      ║
║  ✅ Chọn được bất kỳ Bộ/Ngành nào                            ║
║  ✅ Tải file hướng dẫn (.doc)                                ║
║  ✅ Tự động convert .doc → .docx                             ║
║  ✅ Thêm link thủ tục vào cuối file Word                     ║
║  ✅ Ghi rõ tên thủ tục đầy đủ                                ║
║  ✅ Enhanced RAG integration with semantic search            ║
║  ✅ Improved chunking and embedding capabilities             ║
║  ✅ Better error handling and retry mechanisms               ║
╚═══════════════════════════════════════════════════════════════╝

📦 CÀI ĐẶT:
   pip install selenium requests python-docx

   Windows: pip install pywin32
   Linux:   sudo apt-get install libreoffice
   Mac:     brew install libreoffice
""")

    print("📋 DANH SÁCH BỘ/NGÀNH:")
    print("="*70)
    for key, name in sorted(MINISTRIES.items(), key=lambda x: int(x[0])):
        print(f"{key:>2}. {name}")

    print("\n⚙️ CHỌN BỘ/NGÀNH:")
    choice = get_user_input("Nhập số (1-20): ", "1", MINISTRIES.keys())
    ministry_name = MINISTRIES.get(choice)

    if not ministry_name:
        print("❌ Lựa chọn không hợp lệ!")
        print("✅ Sử dụng giá trị mặc định: Bộ Công an (1)")
        ministry_name = MINISTRIES.get("1")
        choice = "1"

    print(f"\n✅ Đã chọn: {ministry_name}")

    print("\n⚙️ CHỌN CHẾ ĐỘ:")
    print("1. TEST - 5 thủ tục (nhanh, để test)")
    print("2. MEDIUM - 20 thủ tục (kiểm tra)")
    print("3. FULL - TẤT CẢ thủ tục (chạy thật)")

    mode = get_user_input("\nChọn (1-3) [Mặc định: 1]: ", "1", ["1", "2", "3"])
    if not mode:
        mode = "1"

    limits = {"1": 5, "2": 20, "3": None}
    max_procs = limits.get(mode, 5)

    headless = (mode == "3")

    print(f"\n🚀 BẮT ĐẦU...\n")

    url = "https://thutuc.dichvucong.gov.vn/p/home/dvc-tthc-thu-tuc-hanh-chinh.html"

    try:
        scraper = EnhancedProcedureScraper(
            download_dir='downloads_ministries',
            max_workers=8,
            headless=headless
        )
        scraper.run(url, ministry_name, max_procs, store_to_vector=True)
    except KeyboardInterrupt:
        print("\n\n⚠️ Đã dừng bởi người dùng")
    except Exception as e:
        print(f"\n❌ Lỗi: {e}")
        import traceback
        traceback.print_exc()


if __name__ == "__main__":
    import sys
    import io
    # Set stdout to handle Unicode characters properly
    if sys.stdout.encoding != 'utf-8':
        sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
    main()