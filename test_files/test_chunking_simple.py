#!/usr/bin/env python3
"""
Simple test script to verify chunking functionality in the RAG system
"""

import sys
import os
import tempfile

# Add the Knowlegd-rag directory to the Python path
sys.path.insert(0, os.path.join(os.path.dirname(__file__), 'Knowlegd-rag'))

def test_chunking_functionality():
    """Test the chunking functionality with sample content"""
    print("🧪 Testing chunking functionality...")
    
    # Import content processor after adding path
    from content_processor import ContentProcessor
    
    # Initialize content processor
    processor = ContentProcessor()
    
    # Sample Vietnamese administrative procedure content
    sample_content = """
Đây là nội dung thủ tục hành chính mẫu để kiểm tra chức năng chia chunk.
Thủ tục này có mã số TTHC-12345.
Tên thủ tục: Đăng ký kinh doanh.
Cơ quan thực hiện: Sở Kế hoạch và Đầu tư.
Thời hạn giải quyết: 03 ngày làm việc.
Phí, lệ phí: 500.000 đồng.
Hồ sơ bao gồm: Giấy đề nghị, bản sao CMND, giấy tờ chứng minh nơi đặt trụ sở.
Trình tự thực hiện: 
1. Nộp hồ sơ tại bộ phận một cửa.
2. Chờ xử lý và bổ sung nếu cần.
3. Nhận kết quả.

Đây là đoạn văn bản dài hơn để kiểm tra chức năng chia nhỏ thành các phần hợp lý.
Quy trình này áp dụng cho tất cả các tổ chức, cá nhân muốn thành lập doanh nghiệp.
Lĩnh vực: Đầu tư - Kinh doanh.
Các bước thực hiện chi tiết sẽ được nêu rõ trong các điều khoản sau.
Các thành phần hồ sơ cần có đã được liệt kê đầy đủ.
Người thực hiện có trách nhiệm kiểm tra tính hợp lệ của hồ sơ.
Sau khi tiếp nhận, cơ quan có thẩm quyền sẽ xử lý theo đúng quy định.
Thời gian giải quyết có thể thay đổi tùy theo loại hình doanh nghiệp.
Các trường hợp đặc biệt sẽ được xử lý theo quy định riêng.
    """
    
    # Test basic chunking
    print("📝 Testing basic chunking...")
    result = processor.chunk_content_improved(sample_content)
    
    if result.success:
        print(f"✅ Chunking successful: {len(result.chunks)} chunks created")
        for i, chunk in enumerate(result.chunks):
            print(f"   Chunk {i+1}: {len(chunk)} characters")
            # Print first 100 characters of each chunk
            print(f"     Preview: {chunk[:100]}{'...' if len(chunk) > 100 else ''}")
    else:
        print(f"❌ Chunking failed: {result.error}")
        return False
    
    # Test chunking with metadata
    print("\n🏷️ Testing chunking with metadata...")
    metadata = {
        'procedure_code': 'TTHC-12345',
        'title': 'Đăng ký kinh doanh',
        'agency': 'Sở Kế hoạch và Đầu tư',
        'processing_time': '03 ngày làm việc',
        'fee': '500.000 đồng'
    }
    
    result_with_meta = processor.chunk_content_improved(sample_content, metadata)
    
    if result_with_meta.success:
        print(f"✅ Chunking with metadata successful: {len(result_with_meta.chunks)} chunks")
        # Check if metadata is included in first chunk
        if result_with_meta.chunks:
            first_chunk = result_with_meta.chunks[0]
            print(f"   First chunk has metadata header: {'Mã thủ tục:' in first_chunk}")
    else:
        print(f"❌ Chunking with metadata failed: {result_with_meta.error}")
        return False
    
    # Test content extraction
    print("\n🔍 Testing content extraction...")
    extracted_info = processor.extract_structural_info_from_content(sample_content)
    import json
    print(f"   Extracted info: {json.dumps(extracted_info, ensure_ascii=False, indent=2)}")
    
    # Test content preprocessing
    print("\n⚙️ Testing content preprocessing...")
    preprocessed = processor.preprocess_content_for_embedding(sample_content, metadata)
    print(f"   Preprocessed content length: {len(preprocessed)} characters")
    print(f"   Contains structured info: {'TIÊU ĐỀ THỦ TỤC:' in preprocessed}")
    
    print("\n🎉 Chunking functionality test completed successfully!")
    return True

def test_sample_docx_creation():
    """Create a sample docx file for testing"""
    print("\n📄 Creating sample DOCX file for testing...")
    
    try:
        from docx import Document
        
        # Create a temporary docx file
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            doc = Document()
            
            # Add content similar to what we'd find in real procedure documents
            doc.add_heading('THỦ TỤC: Đăng ký kinh doanh', 0)
            doc.add_paragraph('Mã thủ tục: TTHC-12345')
            doc.add_paragraph('Cơ quan thực hiện: Sở Kế hoạch và Đầu tư')
            doc.add_paragraph('Thời hạn giải quyết: 03 ngày làm việc')
            doc.add_paragraph('Phí, lệ phí: 500.000 đồng')
            
            doc.add_heading('Thành phần hồ sơ', 1)
            doc.add_paragraph('1. Giấy đề nghị đăng ký doanh nghiệp')
            doc.add_paragraph('2. Danh sách thành viên/cổ đông sáng lập')
            doc.add_paragraph('3. Bản sao hợp lệ giấy tờ tùy thân của người đại diện')
            
            doc.add_heading('Trình tự thực hiện', 1)
            doc.add_paragraph('Bước 1: Nộp hồ sơ tại bộ phận một cửa')
            doc.add_paragraph('Bước 2: Chờ xử lý và bổ sung nếu cần')
            doc.add_paragraph('Bước 3: Nhận kết quả')
            
            # Add a table (common in procedure docs)
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'STT'
            hdr_cells[1].text = 'Nội dung'
            hdr_cells[2].text = 'Thời hạn'
            
            row_cells = table.add_row().cells
            row_cells[0].text = '1'
            row_cells[1].text = 'Tiếp nhận hồ sơ'
            row_cells[2].text = 'Ngay trong ngày'
            
            row_cells = table.add_row().cells
            row_cells[0].text = '2'
            row_cells[1].text = 'Xử lý hồ sơ'
            row_cells[2].text = '03 ngày làm việc'
            
            doc.save(tmp.name)
            print(f"✅ Sample DOCX file created: {tmp.name}")
            
            # Test extracting content from this file
            print("\n🔍 Testing content extraction from DOCX...")
            from content_processor import ContentProcessor
            processor = ContentProcessor()
            extraction_result = processor.extract_content_from_docx_path(tmp.name)
            
            if extraction_result.success:
                print(f"✅ Content extracted successfully: {len(extraction_result.content)} characters")
                print(f"   Metadata extracted: {bool(extraction_result.metadata)}")
                
                # Test chunking the extracted content
                chunk_result = processor.chunk_content_improved(extraction_result.content, extraction_result.metadata)
                if chunk_result.success:
                    print(f"✅ Content chunked successfully: {len(chunk_result.chunks)} chunks")
                    return tmp.name
                else:
                    print(f"❌ Chunking failed: {chunk_result.error}")
                    return None
            else:
                print(f"❌ Content extraction failed: {extraction_result.error}")
                return None
                
    except Exception as e:
        print(f"❌ Error creating sample DOCX: {e}")
        import traceback
        traceback.print_exc()
        return None

def main():
    """Main test function"""
    print("🔧 Testing RAG System Chunking and Processing")
    print("="*60)
    
    # Test basic functionality
    success = test_chunking_functionality()
    
    if success:
        # Test with actual DOCX file
        docx_path = test_sample_docx_creation()
        
        if docx_path:
            print(f"\n✅ All tests passed! Sample file created at: {docx_path}")
            return True
        else:
            print("\n❌ DOCX test failed")
            return False
    else:
        print("\n❌ Basic functionality test failed")
        return False

if __name__ == "__main__":
    success = main()
    if success:
        print("\n🎉 All RAG system tests completed successfully!")
    else:
        print("\n❌ Some tests failed!")
        sys.exit(1)